"""Issue F3 (teal highlight snap) + F5 (integer char-scaling)."""
from __future__ import annotations

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn

from pdf2docx_plus._vendored.pdf2docx.common.docx import (
    set_char_scaling,
    set_char_shading,
    _nearest_highlight_color,
    _hsv_saturation,
)


def _run():
    return Document().add_paragraph().add_run("x")


# -- F5: integer ST_TextScale ----------------------------------------

def test_char_scaling_value_is_integer():
    r = _run()
    set_char_scaling(r, 0.9666666)
    w = r._r.find(qn("w:rPr")).find(qn("w:w"))
    val = w.get(qn("w:val"))
    assert val == "97", val
    assert "." not in val  # schema requires integer


# -- F3: teal / saturated highlight snap -----------------------------

def test_teal_snaps_to_teal_highlight():
    r = _run()
    set_char_shading(r, 0x007F7F)  # teal editorial ink
    assert r.font.highlight_color == WD_COLOR_INDEX.TEAL
    # must NOT fall back to a <w:shd> box
    assert r._r.find(qn("w:rPr")).find(qn("w:shd")) is None


def test_exact_green_still_bright_green():
    r = _run()
    set_char_shading(r, 0x00FF00)
    assert r.font.highlight_color == WD_COLOR_INDEX.BRIGHT_GREEN


def test_low_saturation_fill_uses_shading_not_highlight():
    r = _run()
    set_char_shading(r, 0x4A4A4A)  # near-grey
    assert _nearest_highlight_color(0x4A4A4A) is None
    assert r.font.highlight_color is None
    assert r._r.find(qn("w:rPr")).find(qn("w:shd")) is not None


def test_saturation_helper():
    assert _hsv_saturation(0x000000) == 0.0
    assert _hsv_saturation(0xFF0000) == 1.0
    assert 0.0 < _hsv_saturation(0x007F7F) <= 1.0
