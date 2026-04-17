"""Tests for ``fit_oversized_tables`` - clamps table indent and
column widths so the table cannot extend past its section's
content area and be clipped by the renderer.
"""

from __future__ import annotations

import importlib.util
from pathlib import Path

import pytest
from docx import Document  # type: ignore
from docx.oxml import OxmlElement  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from docx.shared import Twips  # type: ignore

_SPEC = importlib.util.spec_from_file_location(
    "_table_fit_under_test",
    Path(__file__).resolve().parent.parent / "pdf2docx_plus" / "emit" / "table_fit.py",
)
assert _SPEC and _SPEC.loader
_MOD = importlib.util.module_from_spec(_SPEC)
_SPEC.loader.exec_module(_MOD)
fit_oversized_tables = _MOD.fit_oversized_tables
align_tblgrid_to_cells = _MOD.align_tblgrid_to_cells


def _set_page(doc, *, width_twips=11906, left_twips=720, right_twips=720) -> None:
    """Configure the default section's page size and margins."""
    s = doc.sections[0]
    s.page_width = Twips(width_twips)
    s.left_margin = Twips(left_twips)
    s.right_margin = Twips(right_twips)


def _add_table(doc, col_widths: list[int], tbl_ind: int | None = None) -> "any":
    """Append a fixed-layout table with the given grid column widths
    and optional ``<w:tblInd>`` (in twips). Returns the ``<w:tbl>``
    element for further inspection."""
    tbl = doc.add_table(rows=1, cols=len(col_widths))
    tbl_el = tbl._element
    # clear the auto-generated grid and rebuild with our widths
    grid = tbl_el.find(qn("w:tblGrid"))
    for gc in list(grid.findall(qn("w:gridCol"))):
        grid.remove(gc)
    for w in col_widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        grid.append(gc)
    # set tblInd
    if tbl_ind is not None:
        tblPr = tbl_el.find(qn("w:tblPr"))
        old = tblPr.find(qn("w:tblInd"))
        if old is not None:
            tblPr.remove(old)
        ind = OxmlElement("w:tblInd")
        ind.set(qn("w:w"), str(tbl_ind))
        ind.set(qn("w:type"), "dxa")
        tblPr.append(ind)
    # set cell widths to match grid
    for i, cell in enumerate(tbl.rows[0].cells):
        tcPr = cell._tc.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = OxmlElement("w:tcPr")
            cell._tc.insert(0, tcPr)
        tcW = tcPr.find(qn("w:tcW"))
        if tcW is None:
            tcW = OxmlElement("w:tcW")
            tcPr.append(tcW)
        tcW.set(qn("w:w"), str(col_widths[i]))
        tcW.set(qn("w:type"), "dxa")
    return tbl_el


def _grid_widths(tbl_el) -> list[int]:
    grid = tbl_el.find(qn("w:tblGrid"))
    return [int(gc.get(qn("w:w"))) for gc in grid.findall(qn("w:gridCol"))]


def _tbl_ind(tbl_el) -> int:
    tblPr = tbl_el.find(qn("w:tblPr"))
    ind = tblPr.find(qn("w:tblInd"))
    return int(ind.get(qn("w:w"))) if ind is not None else 0


@pytest.mark.unit
def test_no_change_when_table_fits() -> None:
    doc = Document()
    _set_page(doc, width_twips=11906, left_twips=720, right_twips=720)
    # content width = 10466, table = 4000, indent = 100 -> fits
    tbl = _add_table(doc, col_widths=[2000, 2000], tbl_ind=100)
    fixed = fit_oversized_tables(doc)
    assert fixed == 0
    assert _tbl_ind(tbl) == 100
    assert _grid_widths(tbl) == [2000, 2000]


@pytest.mark.unit
def test_reduces_indent_when_it_alone_causes_overflow() -> None:
    """Table is narrow enough; only the indent pushes it past the
    right edge. The pass should reduce the indent only."""
    doc = Document()
    _set_page(doc, width_twips=11906, left_twips=720, right_twips=720)
    # content = 10466, cols total = 4000. tblInd = 8662 -> ends at 12662,
    # 2196 past content. Expected: indent = 10466 - 4000 = 6466.
    tbl = _add_table(doc, col_widths=[2000, 2000], tbl_ind=8662)
    fixed = fit_oversized_tables(doc)
    assert fixed == 1
    assert _tbl_ind(tbl) == 6466
    assert _grid_widths(tbl) == [2000, 2000]


@pytest.mark.unit
def test_scales_widths_when_table_alone_exceeds_content() -> None:
    """A table wider than the whole content area gets proportionally
    scaled after the indent is zeroed."""
    doc = Document()
    _set_page(doc, width_twips=11906, left_twips=720, right_twips=720)
    # content = 10466, cols = 5292 + 5292 = 10584 > 10466, tblInd = 8662
    tbl = _add_table(doc, col_widths=[5292, 5292], tbl_ind=8662)
    fixed = fit_oversized_tables(doc)
    assert fixed == 1
    assert _tbl_ind(tbl) == 0
    new_widths = _grid_widths(tbl)
    assert sum(new_widths) <= 10466
    # each column should be scaled by the same ratio (equal within 1)
    assert abs(new_widths[0] - new_widths[1]) <= 1


@pytest.mark.unit
def test_scales_cell_widths_in_lockstep_with_grid() -> None:
    """Every <w:tcW> must be scaled by the same ratio as the grid so
    that fixed-layout tables render correctly."""
    doc = Document()
    _set_page(doc, width_twips=11906, left_twips=720, right_twips=720)
    tbl = _add_table(doc, col_widths=[6000, 6000], tbl_ind=0)
    original_total = 12000
    fit_oversized_tables(doc)
    # check every cell's tcW is scaled
    tcWs = []
    for tc in tbl.iter(qn("w:tc")):
        tcPr = tc.find(qn("w:tcPr"))
        tcW = tcPr.find(qn("w:tcW"))
        tcWs.append(int(tcW.get(qn("w:w"))))
    new_widths = _grid_widths(tbl)
    # cell widths sum to approximately the grid total
    assert abs(sum(tcWs) - sum(new_widths)) <= 2


@pytest.mark.unit
def test_preserves_source_right_alignment_when_possible() -> None:
    """When the table fits with some indent, the pass keeps the
    table right-aligned against the content edge (max feasible
    indent) rather than collapsing to the left margin."""
    doc = Document()
    _set_page(doc, width_twips=11906, left_twips=720, right_twips=720)
    # content = 10466, table = 4000 -> max indent that keeps it
    # on-page is 6466. We ask for a bigger indent (8000).
    tbl = _add_table(doc, col_widths=[2000, 2000], tbl_ind=8000)
    fit_oversized_tables(doc)
    assert _tbl_ind(tbl) == 6466


@pytest.mark.unit
def test_tolerates_float_string_indent() -> None:
    """Upstream sometimes emits ``w:w="8662.0"`` on ``<w:tblInd>``.
    The pass must not crash."""
    doc = Document()
    _set_page(doc, width_twips=11906, left_twips=720, right_twips=720)
    tbl = _add_table(doc, col_widths=[2000, 2000])
    tblPr = tbl.find(qn("w:tblPr"))
    ind = OxmlElement("w:tblInd")
    ind.set(qn("w:w"), "8662.0")
    ind.set(qn("w:type"), "dxa")
    # replace any existing indent
    old = tblPr.find(qn("w:tblInd"))
    if old is not None:
        tblPr.remove(old)
    tblPr.append(ind)
    fixed = fit_oversized_tables(doc)
    assert fixed == 1


@pytest.mark.unit
def test_skips_tables_without_grid() -> None:
    """A malformed or missing ``<w:tblGrid>`` is a no-op."""
    doc = Document()
    _set_page(doc, width_twips=11906, left_twips=720, right_twips=720)
    tbl = doc.add_table(rows=1, cols=1)
    grid = tbl._element.find(qn("w:tblGrid"))
    tbl._element.remove(grid)
    fixed = fit_oversized_tables(doc)
    assert fixed == 0


# -- align_tblgrid_to_cells ----------------------------------------------------


def _set_row_cell_widths(tbl, row_idx: int, widths: list[int]) -> None:
    """Overwrite each cell's ``<w:tcW>`` on the given row."""
    row = tbl.rows[row_idx]
    for i, cell in enumerate(row.cells):
        tcPr = cell._tc.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = OxmlElement("w:tcPr")
            cell._tc.insert(0, tcPr)
        tcW = tcPr.find(qn("w:tcW"))
        if tcW is None:
            tcW = OxmlElement("w:tcW")
            tcPr.append(tcW)
        tcW.set(qn("w:w"), str(widths[i]))
        tcW.set(qn("w:type"), "dxa")


@pytest.mark.unit
def test_align_grid_rewrites_uniform_grid_from_cells() -> None:
    """Upstream's signature: grid [4723,4723,4723] but cells
    [1494,4644,8002] - grid should be rewritten to match cells."""
    doc = Document()
    # build 3-col table with equal grid but unequal cell widths
    tbl = _add_table(doc, col_widths=[4723, 4723, 4723], tbl_ind=0)
    tbl_proxy = doc.tables[-1]
    _set_row_cell_widths(tbl_proxy, 0, [1494, 4644, 8002])

    rewritten = align_tblgrid_to_cells(doc)

    assert rewritten == 1
    assert _grid_widths(tbl) == [1494, 4644, 8002]


@pytest.mark.unit
def test_align_grid_noop_when_distribution_already_matches() -> None:
    doc = Document()
    tbl = _add_table(doc, col_widths=[1500, 4600, 8000], tbl_ind=0)
    tbl_proxy = doc.tables[-1]
    _set_row_cell_widths(tbl_proxy, 0, [1500, 4600, 8000])
    rewritten = align_tblgrid_to_cells(doc)
    assert rewritten == 0


@pytest.mark.unit
def test_align_grid_ignores_rows_with_gridspan() -> None:
    """A fully-spanned row must not be used as the canonical source -
    its cell widths represent a merge, not the grid structure."""
    doc = Document()
    tbl = _add_table(doc, col_widths=[4723, 4723, 4723], tbl_ind=0)
    tbl_proxy = doc.tables[-1]
    # row 0: normal cells
    _set_row_cell_widths(tbl_proxy, 0, [1494, 4644, 8002])

    # add a second row with a gridSpan=3 merged cell
    tr = OxmlElement("w:tr")
    tc = OxmlElement("w:tc")
    tcPr = OxmlElement("w:tcPr")
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), "14166")
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)
    span = OxmlElement("w:gridSpan")
    span.set(qn("w:val"), "3")
    tcPr.append(span)
    tc.append(tcPr)
    tr.append(tc)
    tbl.append(tr)

    rewritten = align_tblgrid_to_cells(doc)
    assert rewritten == 1
    # grid must come from the canonical unspanned row, not the merged one
    assert _grid_widths(tbl) == [1494, 4644, 8002]


@pytest.mark.unit
def test_align_grid_skips_when_no_unspanned_row() -> None:
    """A table whose every row is merged has no source of truth; skip."""
    doc = Document()
    tbl = _add_table(doc, col_widths=[4000, 4000], tbl_ind=0)
    tbl_proxy = doc.tables[-1]
    # mark the only row as fully spanned
    row_cells = tbl.find(qn("w:tr")).findall(qn("w:tc"))
    tcPr = row_cells[0].find(qn("w:tcPr"))
    span = OxmlElement("w:gridSpan")
    span.set(qn("w:val"), "2")
    tcPr.append(span)

    rewritten = align_tblgrid_to_cells(doc)
    assert rewritten == 0


@pytest.mark.unit
def test_align_grid_skips_when_cell_count_differs() -> None:
    """Row cell count differs from grid col count (rare, but
    possible after content manipulation). Do not touch the grid."""
    doc = Document()
    tbl = _add_table(doc, col_widths=[2000, 2000, 2000], tbl_ind=0)
    tbl_proxy = doc.tables[-1]
    # remove the third cell from row 0
    row_cells = tbl.find(qn("w:tr")).findall(qn("w:tc"))
    row_cells[-1].getparent().remove(row_cells[-1])
    # row now has 2 cells; grid has 3
    rewritten = align_tblgrid_to_cells(doc)
    assert rewritten == 0
