"""Unit tests for `flatten_per_page_sections`.

Regression coverage for page-count inflation: upstream emits one
`<w:sectPr>` per source PDF page with a default (`nextPage`) break
type; when font substitution causes content to overflow a tight
per-page margin, the next section's hard page break still fires,
costing a full page per overflow.
"""

from __future__ import annotations

import importlib.util
from pathlib import Path

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Avoid pdf2docx_plus.__init__ which pulls in fitz.
_SPEC = importlib.util.spec_from_file_location(
    "_sections_under_test",
    Path(__file__).resolve().parent.parent
    / "pdf2docx_plus"
    / "emit"
    / "sections.py",
)
assert _SPEC and _SPEC.loader
_MOD = importlib.util.module_from_spec(_SPEC)
_SPEC.loader.exec_module(_MOD)
flatten_per_page_sections = _MOD.flatten_per_page_sections


def _add_mid_doc_section(
    doc, *, with_header_ref: bool = False, pg_w: str = "12240", pg_h: str = "15840"
):
    """Add a paragraph that ends a section (mid-doc sectPr).

    Defaults match python-docx's default final sectPr (US Letter) so
    page-size uniformity holds across the test document.
    """
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    sect = OxmlElement("w:sectPr")
    pgSz = OxmlElement("w:pgSz")
    pgSz.set(qn("w:w"), pg_w)
    pgSz.set(qn("w:h"), pg_h)
    sect.append(pgSz)
    if with_header_ref:
        ref = OxmlElement("w:headerReference")
        ref.set(qn("w:type"), "default")
        ref.set(qn("r:id"), "rId1")
        sect.insert(0, ref)
    pPr.append(sect)
    return sect


@pytest.mark.unit
def test_flatten_converts_default_to_continuous() -> None:
    doc = Document()
    s1 = _add_mid_doc_section(doc)
    s2 = _add_mid_doc_section(doc)
    converted = flatten_per_page_sections(doc)
    assert converted == 2
    for sect in (s1, s2):
        type_el = sect.find(qn("w:type"))
        assert type_el is not None
        assert type_el.get(qn("w:val")) == "continuous"


@pytest.mark.unit
def test_flatten_skipped_when_section_has_header_ref() -> None:
    doc = Document()
    _add_mid_doc_section(doc, with_header_ref=True)
    _add_mid_doc_section(doc)
    converted = flatten_per_page_sections(doc)
    assert converted == 0


@pytest.mark.unit
def test_flatten_skipped_on_mixed_page_sizes() -> None:
    doc = Document()
    _add_mid_doc_section(doc)  # default Letter
    _add_mid_doc_section(doc, pg_w="11906", pg_h="16838")  # A4 — different size
    converted = flatten_per_page_sections(doc)
    assert converted == 0


@pytest.mark.unit
def test_flatten_idempotent() -> None:
    doc = Document()
    _add_mid_doc_section(doc)
    _add_mid_doc_section(doc)
    flatten_per_page_sections(doc)
    second = flatten_per_page_sections(doc)
    assert second == 0


@pytest.mark.unit
def test_flatten_no_op_on_single_section_doc() -> None:
    # default Document() has only the final sectPr — nothing to flatten
    doc = Document()
    converted = flatten_per_page_sections(doc)
    assert converted == 0
