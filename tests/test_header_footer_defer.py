"""``extract_headers_footers`` must defer "N Last update: <date>" footer
lines to ``promote_page_numbers_to_footer``.

Those lines embed a page number. If the text extractor lifts a
representative copy into the footer verbatim, the page number is frozen
("1") on every page. Leaving them in the body lets the dedicated
page-footer pass install a live ``PAGE`` field instead.
"""

from __future__ import annotations

import pytest
from docx import Document  # type: ignore
from docx.oxml import OxmlElement  # type: ignore
from docx.oxml.ns import qn  # type: ignore

from pdf2docx_plus.emit.headers_footers import extract_headers_footers
from pdf2docx_plus.emit.page_footer import promote_page_numbers_to_footer
from pdf2docx_plus.layout.hf_detect import HeaderFooter


def _append_sect_break(doc) -> None:
    body = doc.element.body
    final_sect = body.find(qn("w:sectPr"))
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pPr.append(OxmlElement("w:sectPr"))
    p.append(pPr)
    assert final_sect is not None
    final_sect.addprevious(p)


@pytest.mark.unit
def test_last_update_footer_deferred_then_promoted_with_page_field() -> None:
    doc = Document()
    for i in range(1, 4):
        doc.add_paragraph(f"body content for page {i}")
        doc.add_paragraph(f"{i} Last update: 29 November 2024")
        if i < 3:
            _append_sect_break(doc)

    detected = [
        HeaderFooter(
            text="# Last update: # November #",
            bbox=(0.0, 800.0, 300.0, 820.0),
            is_header=False,
            page_ids=(0, 1, 2),
        )
    ]

    # extract must NOT pull the "Last update" line into the footer
    moved = extract_headers_footers(doc, detected)
    assert moved == 0
    assert any("Last update" in p.text for p in doc.paragraphs)

    # the dedicated pass promotes it with a live PAGE field
    promoted = promote_page_numbers_to_footer(doc)
    assert promoted >= 3
    assert not any("Last update" in p.text for p in doc.paragraphs)
    footer_xml = doc.sections[0].footer._element.xml
    assert "PAGE" in footer_xml
    assert "Last update: 29 November 2024" in footer_xml
