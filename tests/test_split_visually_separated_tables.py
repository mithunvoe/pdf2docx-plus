"""Issue P-3: split mega-tables fused from stacked logical tables.

Upstream's stream-table promoter aggregates vertically-adjacent tables
that share a column grid into a single TableBlock.  Fund-prospectus
layouts routinely have multiple fee tables on the same page with the
same column grid; the converter therefore emits them as a single
mega-table.  This breaks any downstream diff that aligns tables by
index.

`split_visually_separated_tables` detects internal header-row
repetition and splits the table at those boundaries.  It is safe
relative to ``stitch_cross_page_tables`` because that path drops the
repeated header before merging, so a correctly-stitched FAQ table never
contains an internal header repeat.
"""

from __future__ import annotations

import importlib.util
from pathlib import Path

import pytest
from docx import Document

_SPEC = importlib.util.spec_from_file_location(
    "_tables_cleanup_under_test",
    Path(__file__).resolve().parent.parent
    / "pdf2docx_plus"
    / "emit"
    / "tables_cleanup.py",
)
assert _SPEC and _SPEC.loader
_MOD = importlib.util.module_from_spec(_SPEC)
_SPEC.loader.exec_module(_MOD)
split_visually_separated_tables = _MOD.split_visually_separated_tables


def _make_table(doc, rows: list[list[str]]) -> None:
    n_rows = len(rows)
    n_cols = max(len(r) for r in rows) if rows else 0
    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    for i, row in enumerate(rows):
        for j, text in enumerate(row):
            tbl.cell(i, j).text = text


@pytest.mark.unit
def test_splits_at_internal_header_repeat() -> None:
    """Three stacked logical tables fused into one — split into three."""
    doc = Document()
    _make_table(
        doc,
        [
            ["Fee", "What you pay"],          # logical table 1 header
            ["Subscription", "Up to 6%"],
            ["Redemption", "Up to 3%"],
            ["Fee", "What you pay"],          # logical table 2 header
            ["Management", "0.45%"],
            ["Custody", "0.05%"],
            ["Fee", "What you pay"],          # logical table 3 header
            ["Performance", "Not applicable"],
        ],
    )
    assert len(doc.tables) == 1
    introduced = split_visually_separated_tables(doc)
    assert introduced == 2  # 1 original + 2 new = 3 tables
    assert len(doc.tables) == 3

    # the three resulting tables should each start with the header row
    headers = [
        (
            doc.tables[i].cell(0, 0).text.strip().casefold(),
            doc.tables[i].cell(0, 1).text.strip().casefold(),
        )
        for i in range(3)
    ]
    assert all(h == ("fee", "what you pay") for h in headers)

    # body rows preserved per segment
    assert doc.tables[0].cell(1, 0).text == "Subscription"
    assert doc.tables[1].cell(1, 0).text == "Management"
    assert doc.tables[2].cell(1, 0).text == "Performance"


@pytest.mark.unit
def test_no_split_when_no_internal_header_repeat() -> None:
    doc = Document()
    _make_table(
        doc,
        [
            ["Q", "A"],
            ["What is X?", "It is Y."],
            ["What is Z?", "It is W."],
            ["When does it open?", "Tomorrow."],
        ],
    )
    introduced = split_visually_separated_tables(doc)
    assert introduced == 0
    assert len(doc.tables) == 1


@pytest.mark.unit
def test_no_split_when_header_signature_too_weak() -> None:
    """A single-cell header signature is too weak — refuse to split."""
    doc = Document()
    # row 0 has only one non-empty cell; the default header_repeat_threshold
    # requires at least 2 to be confident this is a header signature.
    _make_table(
        doc,
        [
            ["A", ""],
            ["data1", "data2"],
            ["A", ""],
            ["data3", "data4"],
        ],
    )
    introduced = split_visually_separated_tables(doc)
    assert introduced == 0
    assert len(doc.tables) == 1


@pytest.mark.unit
def test_no_split_when_table_too_short() -> None:
    """A 2-row table can't contain an internal header repeat by definition."""
    doc = Document()
    _make_table(
        doc,
        [
            ["Fee", "What you pay"],
            ["Subscription", "Up to 6%"],
        ],
    )
    introduced = split_visually_separated_tables(doc)
    assert introduced == 0
    assert len(doc.tables) == 1


@pytest.mark.unit
def test_split_preserves_table_chrome() -> None:
    """The clone tables must carry over the table-level XML (tblPr,
    tblGrid)."""
    from docx.oxml.ns import qn

    doc = Document()
    _make_table(
        doc,
        [
            ["Fee", "What you pay"],
            ["Subscription", "Up to 6%"],
            ["Fee", "What you pay"],
            ["Management", "0.45%"],
        ],
    )
    # original tbl will have a tblPr and tblGrid added by python-docx
    introduced = split_visually_separated_tables(doc)
    assert introduced == 1
    for tbl in doc.tables:
        elem = tbl._element
        # tblPr should be carried over
        assert elem.find(qn("w:tblPr")) is not None
        # tblGrid should be carried over (or at least the grid element)
        assert elem.find(qn("w:tblGrid")) is not None


@pytest.mark.unit
def test_case_insensitive_header_match() -> None:
    """Header detection is case-insensitive (FEE vs fee vs Fee)."""
    doc = Document()
    _make_table(
        doc,
        [
            ["Fee", "What you pay"],
            ["Subscription", "Up to 6%"],
            ["FEE", "What you pay"],
            ["Management", "0.45%"],
        ],
    )
    introduced = split_visually_separated_tables(doc)
    assert introduced == 1
    assert len(doc.tables) == 2


@pytest.mark.unit
def test_no_action_on_empty_document() -> None:
    doc = Document()
    introduced = split_visually_separated_tables(doc)
    assert introduced == 0


@pytest.mark.unit
def test_faq_style_stitched_table_is_not_split() -> None:
    """A correctly-stitched FAQ table has exactly one header at row 0
    — there's no internal repeat — so the splitter must leave it alone.

    This is the guarded invariant called out in PDF_FIDELITY_PDF2DOCX_PLAN.md
    §4.3: P-3 must NOT regress legitimate cross-page FAQ merges.
    """
    doc = Document()
    _make_table(
        doc,
        [
            ["Question", "Answer"],
            ["What is X?", "It is the subscription fee."],
            ["What is Y?", "It is the redemption fee."],
            ["What is Z?", "It is the management charge."],
            ["What is W?", "It is the custody cost."],
            ["What is V?", "It is the performance fee."],
        ],
    )
    introduced = split_visually_separated_tables(doc)
    assert introduced == 0
    assert len(doc.tables) == 1
