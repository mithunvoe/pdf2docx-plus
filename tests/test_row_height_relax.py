"""Issue B5: exact row heights clip cell text and inflate page count.

Upstream ``Row.make_docx`` hardcodes ``height_rule = EXACTLY`` so a row
never grows to fit its content. Under font substitution (Liberation/Carlito
replacing the embedded PDF fonts) the content needs slightly more vertical
space than the source, so the bottom of a cell's text is clipped and the
row is forced onto its own page. Relaxing ``exact`` -> ``atLeast`` keeps the
intended minimum height but lets the row grow to fit.
"""
from __future__ import annotations

import importlib.util
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

_SPEC = importlib.util.spec_from_file_location(
    "_tables_cleanup_b5",
    Path(__file__).resolve().parent.parent / "pdf2docx_plus" / "emit" / "tables_cleanup.py",
)
assert _SPEC and _SPEC.loader
_MOD = importlib.util.module_from_spec(_SPEC)
_SPEC.loader.exec_module(_MOD)
relax_exact_row_heights = _MOD.relax_exact_row_heights


def _set_row_height(tbl, ri: int, val: str, rule: str) -> None:
    tr = tbl.rows[ri]._tr
    trPr = tr.get_or_add_trPr()
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), val)
    h.set(qn("w:hRule"), rule)
    trPr.append(h)


def _hrules(tbl) -> list[str | None]:
    out = []
    for r in tbl.rows:
        trPr = r._tr.find(qn("w:trPr"))
        h = trPr.find(qn("w:trHeight")) if trPr is not None else None
        out.append(h.get(qn("w:hRule")) if h is not None else None)
    return out


def test_exact_rule_relaxed_to_atleast():
    doc = Document()
    tbl = doc.add_table(rows=2, cols=2)
    _set_row_height(tbl, 0, "400", "exact")
    _set_row_height(tbl, 1, "400", "exact")

    changed = relax_exact_row_heights(doc)
    assert changed == 2
    assert _hrules(tbl) == ["atLeast", "atLeast"]
    # the minimum height value is preserved
    tr0 = tbl.rows[0]._tr.find(qn("w:trPr")).find(qn("w:trHeight"))
    assert tr0.get(qn("w:val")) == "400"


def test_atleast_and_auto_rows_untouched():
    doc = Document()
    tbl = doc.add_table(rows=2, cols=1)
    _set_row_height(tbl, 0, "300", "atLeast")
    # row 1 left with no explicit trHeight (auto)
    changed = relax_exact_row_heights(doc)
    assert changed == 0
    assert _hrules(tbl) == ["atLeast", None]
