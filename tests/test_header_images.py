"""Tests for the post-emit ``promote_header_images_to_section`` pass.

Upstream emits the repeating per-page letterhead image as an inline (or
floating) body drawing on every source PDF page. Left in the body it

  * overlaps the first body line when emitted as a floating anchor, and
  * keeps an otherwise chrome-only section alive, so the page-number
    paragraph that accompanies it renders as a near-blank page.

The pass under test lifts a repeating lone-image paragraph into the
section header (one stored copy, consecutive sections linked) and
removes the inline copies from the body. A one-off figure that appears
on a single page must stay in the body.
"""

from __future__ import annotations

import importlib.util
import io
from pathlib import Path

import pytest
from docx import Document  # type: ignore
from docx.oxml import OxmlElement  # type: ignore
from docx.oxml.ns import qn  # type: ignore

_SPEC = importlib.util.spec_from_file_location(
    "_header_images_under_test",
    Path(__file__).resolve().parent.parent
    / "pdf2docx_plus"
    / "emit"
    / "header_images.py",
)
assert _SPEC and _SPEC.loader
_MOD = importlib.util.module_from_spec(_SPEC)
_SPEC.loader.exec_module(_MOD)
promote_header_images_to_section = _MOD.promote_header_images_to_section

# 1x1 PNG (same payload used by test_inline_images.py)
_TINY_PNG = (
    b"\x89PNG\r\n\x1a\n"
    b"\x00\x00\x00\rIHDR"
    b"\x00\x00\x00\x01\x00\x00\x00\x01"
    b"\x08\x02\x00\x00\x00"
    b"\x90wS\xde"
    b"\x00\x00\x00\x0cIDAT"
    b"\x08\x99c\xf8\xcf\xc0\x00\x00\x00\x03\x00\x01"
    b"^\xf3*\xc6"
    b"\x00\x00\x00\x00IEND\xaeB`\x82"
)


def _append_sect_break(doc) -> None:
    body = doc.element.body
    final_sect = body.find(qn("w:sectPr"))
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    sectPr = OxmlElement("w:sectPr")
    pPr.append(sectPr)
    p.append(pPr)
    assert final_sect is not None
    final_sect.addprevious(p)


def _add_lone_image(doc):
    p = doc.add_paragraph()
    p.add_run().add_picture(io.BytesIO(_TINY_PNG))
    return p


def _body_drawing_count(doc) -> int:
    return sum(1 for _ in doc.element.body.iter(qn("w:drawing")))


def _any_header_has_drawing(doc) -> bool:
    for s in doc.sections:
        if list(s.header._element.iter(qn("w:drawing"))):
            return True
    return False


def _build_repeating_letterhead(n_sections: int):
    """A doc with a lone letterhead image at the top of every section."""
    doc = Document()
    for i in range(n_sections):
        _add_lone_image(doc)
        doc.add_paragraph(f"Body content for page {i}")
        if i < n_sections - 1:
            _append_sect_break(doc)
    return doc


@pytest.mark.unit
def test_promotes_repeating_letterhead_to_header() -> None:
    doc = _build_repeating_letterhead(3)
    assert _body_drawing_count(doc) == 3
    moved = promote_header_images_to_section(doc)
    assert moved == 3
    assert _body_drawing_count(doc) == 0
    assert _any_header_has_drawing(doc)
    # body text survives
    texts = [p.text for p in doc.paragraphs]
    assert "Body content for page 0" in texts
    assert "Body content for page 2" in texts


@pytest.mark.unit
def test_single_one_off_image_left_in_body() -> None:
    """A figure that appears on a single section must NOT migrate."""
    doc = Document()
    doc.add_paragraph("intro")
    _add_lone_image(doc)
    doc.add_paragraph("after the figure")
    moved = promote_header_images_to_section(doc)
    assert moved == 0
    assert _body_drawing_count(doc) == 1
    assert not _any_header_has_drawing(doc)


@pytest.mark.unit
def test_image_sharing_paragraph_with_text_not_promoted() -> None:
    """An image inline with real text (not a lone-image paragraph) must
    stay in the body even when it repeats across sections."""
    doc = Document()
    for i in range(3):
        p = doc.add_paragraph(f"caption {i} ")
        p.add_run().add_picture(io.BytesIO(_TINY_PNG))
        if i < 2:
            _append_sect_break(doc)
    moved = promote_header_images_to_section(doc)
    assert moved == 0
    assert _body_drawing_count(doc) == 3


@pytest.mark.unit
def test_idempotent() -> None:
    doc = _build_repeating_letterhead(3)
    promote_header_images_to_section(doc)
    again = promote_header_images_to_section(doc)
    assert again == 0
