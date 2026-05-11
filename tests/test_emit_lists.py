"""Tests for list emission post-pass."""

from __future__ import annotations

import pytest
from docx import Document  # type: ignore

from pdf2docx_plus.emit.lists import apply_lists


@pytest.mark.unit
def test_converts_bullet_to_numPr() -> None:
    doc = Document()
    doc.add_paragraph("• first item")
    doc.add_paragraph("• second item")
    count = apply_lists(doc)
    assert count == 2
    # numPr must now be present on both paragraphs
    for p in doc.paragraphs:
        pPr = p._p.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr")
        assert pPr is not None, "pPr missing"
        numPr = pPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr")
        assert numPr is not None, "numPr missing"


@pytest.mark.unit
def test_decimal_list_detected() -> None:
    doc = Document()
    doc.add_paragraph("1. first")
    doc.add_paragraph("2. second")
    doc.add_paragraph("3. third")
    count = apply_lists(doc)
    assert count == 3


@pytest.mark.unit
def test_mixed_content_only_converts_list_paragraphs() -> None:
    """A run of consecutive same-kind markers should convert; a single
    isolated marker (no run-mate) should be left alone since
    single-paragraph "lists" are usually decorative noise rather than
    real lists.
    """
    doc = Document()
    doc.add_paragraph("Introduction text.")
    doc.add_paragraph("• bullet one")
    doc.add_paragraph("• bullet two")
    doc.add_paragraph("More prose.")
    doc.add_paragraph("1. numbered one")
    doc.add_paragraph("2. numbered two")
    count = apply_lists(doc)
    assert count == 4


@pytest.mark.unit
def test_single_marker_not_promoted() -> None:
    """Single isolated bullet/decimal should NOT be promoted - it's
    almost always decorative or a stray reference, not a list."""
    doc = Document()
    doc.add_paragraph("Introduction text.")
    doc.add_paragraph("• one-off")
    doc.add_paragraph("More prose.")
    count = apply_lists(doc)
    assert count == 0


@pytest.mark.unit
def test_strips_marker_from_text() -> None:
    doc = Document()
    doc.add_paragraph("• hello world")
    doc.add_paragraph("• again here")
    apply_lists(doc)
    assert doc.paragraphs[0].text == "hello world"


@pytest.mark.unit
def test_empty_document_safe() -> None:
    doc = Document()
    assert apply_lists(doc) == 0
