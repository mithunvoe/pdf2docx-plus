"""Unit tests for `promote_page_numbers_to_footer`.

Regression coverage for page numbers rendered as inline body text
instead of in the footer — upstream emits ``"N Last update: ..."`` as
a plain paragraph on every page, which leaves page numbers static when
the doc repaginates and repeats the footer line 67 times in the body.
"""

from __future__ import annotations

import importlib.util
from pathlib import Path

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

_SPEC = importlib.util.spec_from_file_location(
    "_page_footer_under_test",
    Path(__file__).resolve().parent.parent
    / "pdf2docx_plus"
    / "emit"
    / "page_footer.py",
)
assert _SPEC and _SPEC.loader
_MOD = importlib.util.module_from_spec(_SPEC)
_SPEC.loader.exec_module(_MOD)
promote_page_numbers_to_footer = _MOD.promote_page_numbers_to_footer


def _add_body_paragraph(doc, text: str) -> None:
    doc.add_paragraph(text)


def _append_sect_break(doc) -> None:
    """Insert a section-break paragraph before the final body sectPr."""
    body = doc.element.body
    final_sect = body.find(qn("w:sectPr"))
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    sectPr = OxmlElement("w:sectPr")
    pPr.append(sectPr)
    p.append(pPr)
    assert final_sect is not None
    final_sect.addprevious(p)


@pytest.mark.unit
def test_promotes_merged_page_number_and_last_update() -> None:
    doc = Document()
    for i in range(1, 4):
        _add_body_paragraph(doc, f"{i} Last update: 2 October 2024")
    promoted = promote_page_numbers_to_footer(doc)
    assert promoted == 3
    body_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Last update" not in body_text


@pytest.mark.unit
def test_promotes_bare_page_number_before_footer_line() -> None:
    doc = Document()
    _add_body_paragraph(doc, "real body content")
    _add_body_paragraph(doc, "5")
    _add_body_paragraph(doc, "Last update: 2 October 2024")
    promoted = promote_page_numbers_to_footer(doc)
    assert promoted == 2  # bare "5" + "Last update: ..."
    # body should retain "real body content"
    body_texts = [p.text for p in doc.paragraphs]
    assert "real body content" in body_texts
    assert "5" not in body_texts
    assert not any("Last update" in t for t in body_texts)


@pytest.mark.unit
def test_footer_has_page_field() -> None:
    doc = Document()
    _add_body_paragraph(doc, "1 Last update: 2 October 2024")
    promote_page_numbers_to_footer(doc)
    footer = doc.sections[0].footer
    ftr_xml = footer._element.xml  # type: ignore[attr-defined]
    assert "PAGE" in ftr_xml
    assert 'w:fldCharType="begin"' in ftr_xml
    assert 'w:fldCharType="end"' in ftr_xml
    assert "Last update: 2 October 2024" in ftr_xml


@pytest.mark.unit
def test_no_op_when_no_footer_pattern_present() -> None:
    doc = Document()
    _add_body_paragraph(doc, "just body text")
    _add_body_paragraph(doc, "42")  # isolated digit, no footer context
    promoted = promote_page_numbers_to_footer(doc)
    assert promoted == 0
    # digit paragraph stays because it doesn't sit next to a footer line
    assert any(p.text == "42" for p in doc.paragraphs)


@pytest.mark.unit
def test_idempotent() -> None:
    doc = Document()
    _add_body_paragraph(doc, "1 Last update: 2 October 2024")
    promote_page_numbers_to_footer(doc)
    second = promote_page_numbers_to_footer(doc)
    assert second == 0


@pytest.mark.unit
def test_promotes_bare_monotonic_page_number_sequence() -> None:
    """First Sentier-style: bare ``"1", "2", ..., "N"`` sprinkled
    one per source page, with no ``Last update:`` line."""
    doc = Document()
    for i in range(1, 7):
        _add_body_paragraph(doc, f"body content of page {i}")
        _add_body_paragraph(doc, str(i))
    promoted = promote_page_numbers_to_footer(doc)
    assert promoted == 6
    body_text = [p.text for p in doc.paragraphs]
    for i in range(1, 7):
        assert str(i) not in body_text
    assert all(f"body content of page {i}" in body_text for i in range(1, 7))
    # Bare-digit path leaves upstream's footer alone — installing a new
    # footer in the tight per-page sections re-inflates the page count.


@pytest.mark.unit
def test_promotes_sequence_with_small_gaps() -> None:
    """Upstream sometimes drops a page number on a full-bleed image
    page; a gap of 2 should still be tolerated."""
    doc = Document()
    for v in (1, 2, 3, 4, 5, 7, 8, 10, 11):
        _add_body_paragraph(doc, f"content_{v}")
        _add_body_paragraph(doc, str(v))
    promoted = promote_page_numbers_to_footer(doc)
    assert promoted == 9
    body_text = [p.text for p in doc.paragraphs]
    for v in (1, 2, 3, 4, 5, 7, 8, 10, 11):
        assert str(v) not in body_text


@pytest.mark.unit
def test_skips_short_monotonic_run() -> None:
    """Fewer than five digits is not enough evidence."""
    doc = Document()
    _add_body_paragraph(doc, "body")
    _add_body_paragraph(doc, "1")
    _add_body_paragraph(doc, "2")
    _add_body_paragraph(doc, "3")
    _add_body_paragraph(doc, "body")
    promoted = promote_page_numbers_to_footer(doc)
    assert promoted == 0


@pytest.mark.unit
def test_skips_when_digits_look_like_data_values() -> None:
    """Scattered digit values (75, 100, 42) are not a page-number run.

    The heuristic requires a monotonic step-1 sequence that starts
    at 1-3 and covers the majority of bare-digit paragraphs.
    """
    doc = Document()
    _add_body_paragraph(doc, "Table values:")
    for val in ("75", "100", "42", "17", "5"):
        _add_body_paragraph(doc, val)
    promoted = promote_page_numbers_to_footer(doc)
    assert promoted == 0
    assert all(p.text in {"75", "100", "42", "17", "5", "Table values:"} for p in doc.paragraphs)


@pytest.mark.unit
def test_sparse_page_run_ignored_when_mixed_with_data() -> None:
    """A short page-run (1,2,3) interleaved with many non-monotonic digits
    should not trigger promotion — the digits are probably data."""
    doc = Document()
    _add_body_paragraph(doc, "1")
    _add_body_paragraph(doc, "2")
    _add_body_paragraph(doc, "3")
    # 7 unrelated digit paragraphs - run is now <50% of bare digits
    for v in ("75", "100", "42", "17", "9", "88", "6"):
        _add_body_paragraph(doc, v)
    promoted = promote_page_numbers_to_footer(doc)
    assert promoted == 0


# -- per-section trailing page numbers (decorated / short sequences) -------


def _build_per_section_trailing(values: list[str]) -> "Document":
    """One section per value; each section ends with a page-number line."""
    doc = Document()
    for i, v in enumerate(values):
        _add_body_paragraph(doc, f"real body content for page {i}")
        _add_body_paragraph(doc, v)
        if i < len(values) - 1:
            _append_sect_break(doc)
    return doc


@pytest.mark.unit
def test_strips_decorated_trailing_page_numbers() -> None:
    """``– 2 –`` style page numbers, one per section, are recognised even
    though they are neither bare digits nor a 'Last update' line."""
    doc = _build_per_section_trailing(["– 2 –", "– 3 –", "– 4 –"])
    promoted = promote_page_numbers_to_footer(doc)
    assert promoted == 3
    body = [p.text for p in doc.paragraphs]
    assert not any("2" in t or "3" in t or "4" in t for t in body if "page" not in t)
    assert all(f"real body content for page {i}" in body for i in range(3))


@pytest.mark.unit
def test_strips_short_per_section_bare_sequence() -> None:
    """A 3-page doc with one bare page number per section is below the
    bare-sequence floor of 5 but is still recognised per-section."""
    doc = _build_per_section_trailing(["1", "2", "3"])
    promoted = promote_page_numbers_to_footer(doc)
    assert promoted == 3
    body = [p.text for p in doc.paragraphs]
    assert "1" not in body and "2" not in body and "3" not in body


@pytest.mark.unit
def test_strips_bracketed_and_word_page_numbers() -> None:
    doc = _build_per_section_trailing(["[2]", "Page 3", "4 of 10"])
    promoted = promote_page_numbers_to_footer(doc)
    assert promoted == 3


@pytest.mark.unit
def test_keeps_non_monotonic_trailing_numbers() -> None:
    """Trailing numbers that do not increase like page numbers are data,
    not pagination — leave them alone."""
    doc = _build_per_section_trailing(["7", "2", "9"])
    promoted = promote_page_numbers_to_footer(doc)
    assert promoted == 0
    body = [p.text for p in doc.paragraphs]
    assert "7" in body and "2" in body and "9" in body


@pytest.mark.unit
def test_keeps_single_trailing_number() -> None:
    """One section ending in a number is not enough evidence."""
    doc = Document()
    _add_body_paragraph(doc, "body")
    _add_body_paragraph(doc, "5")
    promoted = promote_page_numbers_to_footer(doc)
    assert promoted == 0
    assert any(p.text == "5" for p in doc.paragraphs)
