"""Regression tests for list/checkbox formatting fixes.

E1: empty/filled square glyphs (U+25A1 □, U+25A0 ■) must NOT be treated as
bullet markers — in fund/form documents they are checkboxes, and promoting
them to a Word bullet list strips the tickable box semantics.

B2b: list detection must recurse into table cells so bullets trapped inside
pseudo-tables still receive real ``w:numPr``.
"""
from __future__ import annotations

from docx import Document

from pdf2docx_plus.layout.lists import detect_list_block
from pdf2docx_plus.emit.lists import apply_lists


# -- E1 ----------------------------------------------------------------

def test_white_square_is_not_a_bullet():
    assert detect_list_block("□ No") is None


def test_black_square_is_not_a_bullet():
    assert detect_list_block("■ No") is None


def test_checkbox_form_choice_preserved():
    # a checkbox line with binary choice must stay verbatim
    assert detect_list_block("□ Yes □ No") is None


def test_real_bullet_still_detected():
    m = detect_list_block("• first item")
    assert m is not None and m.kind == "bullet"


def test_small_black_square_bullet_still_detected():
    # U+25AA BLACK SMALL SQUARE is a genuine bullet glyph, not a checkbox.
    m = detect_list_block("▪ item one")
    assert m is not None and m.kind == "bullet"


# -- B2b ---------------------------------------------------------------

def test_bullets_inside_table_cell_get_numbering():
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    # first cell paragraph already exists; reuse + append
    cell.paragraphs[0].add_run("• alpha")
    cell.add_paragraph("• bravo")
    cell.add_paragraph("• charlie")

    converted = apply_lists(doc)
    assert converted >= 3, f"expected in-cell bullets promoted, got {converted}"
    # the marker glyph should be stripped from the cell paragraphs
    texts = [p.text for p in cell.paragraphs]
    assert all("•" not in t for t in texts), texts
