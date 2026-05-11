"""Tests for inline image emission patch."""

from __future__ import annotations

import io

import pytest


_TINY_PNG = (
    b"\x89PNG\r\n\x1a\n"
    b"\x00\x00\x00\rIHDR"  # IHDR chunk
    b"\x00\x00\x00\x01\x00\x00\x00\x01"  # 1x1
    b"\x08\x02\x00\x00\x00"  # 8-bit RGB
    b"\x90wS\xde"  # CRC
    b"\x00\x00\x00\x0cIDAT"  # IDAT
    b"\x08\x99c\xf8\xcf\xc0\x00\x00\x00\x03\x00\x01"
    b"^\xf3*\xc6"  # CRC
    b"\x00\x00\x00\x00IEND\xaeB`\x82"  # IEND
)


@pytest.mark.unit
def test_add_image_emits_inline_drawing() -> None:
    """``docx.add_image`` after the fidelity patch should produce a
    ``wp:inline`` rather than a ``wp:anchor`` drawing."""
    # importing pdf2docx_plus installs the fidelity patches
    import pdf2docx_plus  # noqa: F401
    from docx import Document  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    from pdf2docx_plus._vendored.pdf2docx.common.docx import add_image

    doc = Document()
    p = doc.add_paragraph()
    add_image(p, io.BytesIO(_TINY_PNG), 100, 100, 36, 36)

    # look for wp:inline vs wp:anchor in the body
    body = doc.element.body
    has_inline = False
    has_anchor = False
    for drawing in body.iter(qn("w:drawing")):
        for child in drawing:
            tag = child.tag
            if isinstance(tag, str):
                if tag.endswith("}inline"):
                    has_inline = True
                elif tag.endswith("}anchor"):
                    has_anchor = True

    assert has_inline, "expected wp:inline drawing"
    assert not has_anchor, "wp:anchor should be reserved for explicit float images"


@pytest.mark.unit
def test_add_image_drops_unrecognised() -> None:
    """An unrecognised image must not leave an orphan empty run in the paragraph."""
    import pdf2docx_plus  # noqa: F401
    from docx import Document  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    from pdf2docx_plus._vendored.pdf2docx.common.docx import add_image

    doc = Document()
    p = doc.add_paragraph()
    add_image(p, io.BytesIO(b"not a real image"), 0, 0, 10, 10)
    # paragraph should have no <w:drawing>
    drawings = list(p._p.iter(qn("w:drawing")))
    assert drawings == []
