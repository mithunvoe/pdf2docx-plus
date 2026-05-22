"""Unit tests for post-emit table cleanup passes.

Regression coverage for "tables out of nowhere" — upstream's lattice
detector faithfully finds bordered rectangles in the PDF (e.g. empty
checkbox grids, stroke artifacts) but content extraction leaves every
cell blank, so they surface in the DOCX as mysterious empty grids.
"""

from __future__ import annotations

import importlib.util
from pathlib import Path

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Import the module directly to avoid the pdf2docx_plus package __init__,
# which pulls in `fitz` (not needed for these pure-XML cleanup passes).
_SPEC = importlib.util.spec_from_file_location(
    "_tables_cleanup_under_test",
    Path(__file__).resolve().parent.parent / "pdf2docx_plus" / "emit" / "tables_cleanup.py",
)
assert _SPEC and _SPEC.loader
_MOD = importlib.util.module_from_spec(_SPEC)
_SPEC.loader.exec_module(_MOD)
drop_empty_tables = _MOD.drop_empty_tables
trim_empty_table_rows = _MOD.trim_empty_table_rows
unwrap_tiny_tables = _MOD.unwrap_tiny_tables


def _make_doc_with_tables(specs: list[list[list[str]]]):
    """Build a Document with one table per spec. spec = [row[cell_text]]."""
    doc = Document()
    for spec in specs:
        rows = len(spec)
        cols = max(len(r) for r in spec) if spec else 0
        tbl = doc.add_table(rows=rows, cols=cols)
        for ri, row in enumerate(spec):
            for ci, text in enumerate(row):
                tbl.cell(ri, ci).text = text
    return doc


def _add_box_borders(table) -> None:
    """Give every cell inline box borders (start/end/bottom), matching the
    way the upstream lattice emitter draws form-field boxes."""
    for tc in table._tbl.iter(qn("w:tc")):
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = OxmlElement("w:tcPr")
            tc.insert(0, tcPr)
        borders = OxmlElement("w:tcBorders")
        for side in ("start", "end", "bottom"):
            edge = OxmlElement("w:" + side)
            edge.set(qn("w:val"), "single")
            edge.set(qn("w:sz"), "6")
            edge.set(qn("w:color"), "000000")
            borders.append(edge)
        tcPr.append(borders)


@pytest.mark.unit
def test_drop_empty_tables_removes_all_blank() -> None:
    doc = _make_doc_with_tables(
        [
            [["", ""], ["", ""]],  # all blank → drop
            [["hello", ""], ["", ""]],  # one cell has content → keep
        ]
    )
    assert len(doc.tables) == 2
    removed = drop_empty_tables(doc)
    assert removed == 1
    assert len(doc.tables) == 1
    assert doc.tables[0].cell(0, 0).text == "hello"


@pytest.mark.unit
def test_drop_empty_tables_whitespace_is_empty() -> None:
    doc = _make_doc_with_tables([[["   ", "\t"], ["", "\n "]]])
    removed = drop_empty_tables(doc)
    assert removed == 1


@pytest.mark.unit
def test_drop_empty_tables_leaves_populated_sparse_tables() -> None:
    # sparse but not empty — genuine data table with mostly blank cells
    doc = _make_doc_with_tables([[["A", "", ""], ["", "", ""], ["", "", "Z"]]])
    removed = drop_empty_tables(doc)
    assert removed == 0
    assert len(doc.tables) == 1


@pytest.mark.unit
def test_drop_empty_tables_preserves_large_empty_form_grids() -> None:
    """A 14x2 all-empty table is almost certainly a form's checkbox
    continuation grid, not a decoration. Do not drop it."""
    big_empty = [["", ""] for _ in range(14)]
    doc = _make_doc_with_tables([big_empty])
    removed = drop_empty_tables(doc)
    assert removed == 0
    assert len(doc.tables) == 1


@pytest.mark.unit
def test_drop_empty_tables_keeps_bordered_form_box() -> None:
    """A small empty table with real box borders is a form-field grid
    (e.g. the Day/Month/Year or Flat/Room address boxes) and must be
    kept even though it is empty and under the cell-count cap."""
    doc = _make_doc_with_tables([[["", "", "", ""]]])  # 1x4 empty
    _add_box_borders(doc.tables[0])
    removed = drop_empty_tables(doc)
    assert removed == 0
    assert len(doc.tables) == 1


@pytest.mark.unit
def test_drop_empty_tables_keeps_bordered_single_box() -> None:
    """Even a 1x1 empty bordered box is a tick / input field, not noise."""
    doc = _make_doc_with_tables([[[""]]])
    _add_box_borders(doc.tables[0])
    removed = drop_empty_tables(doc)
    assert removed == 0
    assert len(doc.tables) == 1


@pytest.mark.unit
def test_unwrap_tiny_tables_skips_bordered_table() -> None:
    """A bordered single-row table is a genuine lattice table and must
    keep its cell structure rather than being flattened to tab text."""
    doc = _make_doc_with_tables([[["a", "b"]]])
    _add_box_borders(doc.tables[0])
    unwrapped = unwrap_tiny_tables(doc)
    assert unwrapped == 0
    assert len(doc.tables) == 1


@pytest.mark.unit
def test_unwrap_tiny_tables_unwraps_borderless_sprawl() -> None:
    """A borderless single-row table is upstream's stream-detected
    label/value sprawl and is still unwrapped to paragraphs."""
    doc = _make_doc_with_tables([[["a", "b"]]])
    unwrapped = unwrap_tiny_tables(doc)
    assert unwrapped == 1
    assert len(doc.tables) == 0


@pytest.mark.unit
def test_trim_empty_table_rows_trims_sparse_lattice_artifact() -> None:
    """Lattice detector often finds a rectangle around a single piece
    of text plus a few spurious empty cells. The 4-row case with one
    content row is the signature we want to collapse to 1 row."""
    doc = _make_doc_with_tables([[["", ""], ["only", "data"], ["", ""], ["", ""]]])
    tbl = doc.tables[0]
    assert len(tbl.rows) == 4
    trimmed = trim_empty_table_rows(doc)
    assert trimmed == 3
    rows_after = [[c.text for c in r.cells] for r in doc.tables[0].rows]
    assert rows_after == [["only", "data"]]


@pytest.mark.unit
def test_trim_empty_table_rows_preserves_single_row() -> None:
    doc = _make_doc_with_tables([[["", ""]]])
    trimmed = trim_empty_table_rows(doc)
    assert trimmed == 0
    assert len(doc.tables[0].rows) == 1


@pytest.mark.unit
def test_trim_empty_table_rows_preserves_form_checkbox_tables() -> None:
    """Multi-row forms (``Applicable? | Yes | No`` + empty checkbox
    rows) must survive untouched: the empty rows are the form."""
    doc = _make_doc_with_tables(
        [[["Applicable?", "Applicable?"], ["Yes", "No"], ["", ""], ["", ""], ["", ""], ["", ""]]]
    )
    tbl = doc.tables[0]
    assert len(tbl.rows) == 6
    trimmed = trim_empty_table_rows(doc)
    assert trimmed == 0
    assert len(doc.tables[0].rows) == 6


@pytest.mark.unit
def test_trim_empty_table_rows_preserves_multi_content_tables() -> None:
    """A 5-row table with two content rows is a real table; no trimming."""
    doc = _make_doc_with_tables([[["", ""], ["row1", "a"], ["", ""], ["row2", "b"], ["", ""]]])
    trimmed = trim_empty_table_rows(doc)
    assert trimmed == 0
