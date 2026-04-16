"""Tests for the post-emit word-spacing repair pass.

Upstream emits adjacent ``<w:r>`` runs when a sentence wraps across
source-PDF lines; the trailing space at the line break is dropped,
so adjacent runs look like ``["confirms,", "having made..."]`` and
render as ``"confirms,having made..."``. The pass under test
restores the missing space between those runs.

Uses ``importlib.util`` to import the pure-XML pass module without
importing ``pdf2docx_plus`` (which in turn imports ``fitz`` and
requires the dev venv).
"""

from __future__ import annotations

import importlib.util
from pathlib import Path

import pytest
from docx import Document  # type: ignore

_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _load_module():
    path = Path(__file__).resolve().parent.parent / "pdf2docx_plus" / "emit" / "word_spacing.py"
    spec = importlib.util.spec_from_file_location("word_spacing_under_test", path)
    assert spec and spec.loader
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


_mod = _load_module()
repair_wrap_spacing = _mod.repair_wrap_spacing


def _set_para_runs(paragraph, pairs):
    """Replace paragraph's runs with given (text, preserve) pairs."""
    p = paragraph._p
    # strip existing runs
    for r in list(p.findall(_W + "r")):
        p.remove(r)
    for text, preserve in pairs:
        run = paragraph.add_run(text)
        if preserve:
            t = run._r.find(_W + "t")
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")


def _paragraph_text(paragraph):
    return "".join(t.text or "" for t in paragraph._p.iter(_W + "t"))


@pytest.mark.unit
def test_inserts_space_after_comma() -> None:
    doc = Document()
    p = doc.add_paragraph()
    _set_para_runs(p, [("confirms,", False), ("having made enquiries", False)])

    fixed = repair_wrap_spacing(doc)

    assert fixed == 1
    assert _paragraph_text(p) == "confirms, having made enquiries"


@pytest.mark.unit
def test_inserts_space_after_sentence_period() -> None:
    doc = Document()
    p = doc.add_paragraph()
    _set_para_runs(p, [("Sub-Fund.", False), ("The valuation activities", False)])

    fixed = repair_wrap_spacing(doc)

    assert fixed == 1
    assert _paragraph_text(p) == "Sub-Fund. The valuation activities"


@pytest.mark.unit
def test_preserves_single_letter_initials() -> None:
    """``U.S.`` followed by ``Department`` must not gain a space."""
    doc = Document()
    p = doc.add_paragraph()
    _set_para_runs(p, [("U.S.", False), ("Department", False)])

    fixed = repair_wrap_spacing(doc)

    assert fixed == 0
    assert _paragraph_text(p) == "U.S.Department"


@pytest.mark.unit
def test_preserves_hyphenated_line_wrap() -> None:
    """Runs joined with a trailing hyphen must not gain a space."""
    doc = Document()
    p = doc.add_paragraph()
    _set_para_runs(p, [("first-come-", False), ("first-serve", False)])

    fixed = repair_wrap_spacing(doc)

    assert fixed == 0
    assert _paragraph_text(p) == "first-come-first-serve"


@pytest.mark.unit
def test_skips_when_left_already_trails_space() -> None:
    doc = Document()
    p = doc.add_paragraph()
    _set_para_runs(p, [("already, ", True), ("spaced", False)])

    fixed = repair_wrap_spacing(doc)

    assert fixed == 0
    assert _paragraph_text(p) == "already, spaced"


@pytest.mark.unit
def test_skips_when_right_already_leads_space() -> None:
    doc = Document()
    p = doc.add_paragraph()
    _set_para_runs(p, [("first.", False), (" Second", True)])

    fixed = repair_wrap_spacing(doc)

    assert fixed == 0
    assert _paragraph_text(p) == "first. Second"


@pytest.mark.unit
def test_inserts_space_after_closing_paren() -> None:
    doc = Document()
    p = doc.add_paragraph()
    _set_para_runs(p, [("(the Fund)", False), ("The", False)])

    fixed = repair_wrap_spacing(doc)

    assert fixed == 1
    assert _paragraph_text(p) == "(the Fund) The"


@pytest.mark.unit
def test_skips_non_letter_next_run() -> None:
    """``http:`` followed by ``//example.com`` must not split."""
    doc = Document()
    p = doc.add_paragraph()
    _set_para_runs(p, [("http:", False), ("//example.com", False)])

    fixed = repair_wrap_spacing(doc)

    assert fixed == 0
    assert _paragraph_text(p) == "http://example.com"


@pytest.mark.unit
def test_handles_runs_in_table_cells() -> None:
    doc = Document()
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    cell.paragraphs[0].text = ""  # clear
    _set_para_runs(cell.paragraphs[0], [("Item,", False), ("please", False)])

    fixed = repair_wrap_spacing(doc)

    assert fixed == 1
    assert _paragraph_text(cell.paragraphs[0]) == "Item, please"


@pytest.mark.unit
def test_multiple_pairs_in_one_paragraph() -> None:
    doc = Document()
    p = doc.add_paragraph()
    _set_para_runs(
        p,
        [
            ("basis.", False),
            ("For the purposes,", False),
            ("the Trustee", False),
        ],
    )

    fixed = repair_wrap_spacing(doc)

    assert fixed == 2
    assert _paragraph_text(p) == "basis. For the purposes, the Trustee"


@pytest.mark.unit
def test_line_break_between_runs_not_modified() -> None:
    """Runs separated by an explicit ``<w:br/>`` must not gain a
    space - the line break already forces the split."""
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("line one.")
    # append <w:br/> inside the run
    from docx.oxml import OxmlElement  # type: ignore

    br = OxmlElement("w:br")
    r._r.append(br)
    p.add_run("Line two")

    fixed = repair_wrap_spacing(doc)

    assert fixed == 0
