"""Tests for run consolidation."""

from __future__ import annotations

import pytest
from docx import Document  # type: ignore

from pdf2docx_plus.consolidate import consolidate_runs


@pytest.mark.unit
def test_merges_identical_runs() -> None:
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("hello ")
    p.add_run("world ")
    p.add_run("again")
    merged = consolidate_runs(doc)
    assert merged >= 1
    assert "".join(r.text for r in p.runs) == "hello world again"


@pytest.mark.unit
def test_preserves_different_formatting() -> None:
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("normal ")
    bold = p.add_run("bold")
    bold.bold = True
    p.add_run(" normal")
    consolidate_runs(doc)
    # we should still have a separate bold run
    bold_runs = [r for r in p.runs if r.bold]
    assert len(bold_runs) == 1


@pytest.mark.unit
def test_empty_paragraph_safe() -> None:
    doc = Document()
    doc.add_paragraph()
    assert consolidate_runs(doc) == 0
