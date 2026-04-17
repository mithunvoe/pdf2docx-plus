"""Regression tests for fidelity patches (hyperlink OOXML, null-byte strip)."""

from __future__ import annotations

import re

import pytest
from docx import Document  # type: ignore

from pdf2docx_plus import fidelity  # noqa: F401 - import installs patches
from pdf2docx_plus.fidelity.text import sanitize


@pytest.mark.unit
def test_sanitize_strips_invalid_xml_controls() -> None:
    dirty = "hello\x00world\x01\x08test\x0b\x0c\x1e\x1f end"
    assert sanitize(dirty) == "helloworldtest end"


@pytest.mark.unit
def test_sanitize_preserves_valid_whitespace() -> None:
    assert sanitize("a\tb\nc\rd") == "a\tb\nc\rd"


@pytest.mark.unit
def test_sanitize_none_and_empty() -> None:
    assert sanitize(None) == ""
    assert sanitize("") == ""


@pytest.mark.unit
def test_paragraph_add_run_sanitises() -> None:
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("good\x00bad")
    assert r.text == "goodbad"


@pytest.mark.unit
def test_add_hyperlink_is_paragraph_level() -> None:
    """Post-patch, <w:hyperlink> must be a direct child of <w:p>, not <w:r>."""
    from pdf2docx_plus._vendored.pdf2docx.common.docx import add_hyperlink

    doc = Document()
    p = doc.add_paragraph()
    add_hyperlink(p, "https://example.com", "click")

    xml = p._p.xml
    # hyperlink must appear at paragraph level
    assert re.search(r"<w:p[^>]*>.*?<w:hyperlink\b", xml, re.DOTALL)
    # hyperlink must NOT be nested inside a run
    assert not re.search(r"<w:r\b[^>]*>.*?<w:hyperlink", xml, re.DOTALL)
