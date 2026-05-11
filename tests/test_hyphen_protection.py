"""Tests for hyphen protection in wrap-spacing repair."""

from __future__ import annotations

import pytest
from docx import Document  # type: ignore
from docx.oxml import OxmlElement  # type: ignore
from docx.oxml.ns import qn  # type: ignore

from pdf2docx_plus.emit.word_spacing import (
    _HYPHEN_WHITELIST,
    _is_protected_hyphen_pair,
    repair_wrap_spacing,
)


@pytest.mark.unit
def test_whitelist_protects_sub_fund() -> None:
    assert _is_protected_hyphen_pair("Sub-", "Fund")
    assert _is_protected_hyphen_pair("sub-", "fund")
    assert _is_protected_hyphen_pair("the sub-", "fund pays")
    # plural form
    assert _is_protected_hyphen_pair("Sub-", "Funds")


@pytest.mark.unit
def test_whitelist_protects_non_listed() -> None:
    assert _is_protected_hyphen_pair("non-", "listed")
    assert _is_protected_hyphen_pair("the non-", "listed shares")


@pytest.mark.unit
def test_unknown_hyphenated_compound_not_protected() -> None:
    # an arbitrary hyphenated compound the whitelist doesn't know
    assert not _is_protected_hyphen_pair("xyz-", "abc")


@pytest.mark.unit
def test_repair_doesnt_add_space_after_hyphen() -> None:
    """A run ending with '-' must never have a space appended even when
    the next run starts with an uppercase letter."""
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Sub-")
    p.add_run("Fund")
    fixed = repair_wrap_spacing(doc)
    assert fixed == 0
    assert p.text == "Sub-Fund"


@pytest.mark.unit
def test_repair_still_adds_space_after_comma() -> None:
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("alpha,")
    p.add_run("beta")
    fixed = repair_wrap_spacing(doc)
    assert fixed == 1
    assert p.text == "alpha, beta"


@pytest.mark.unit
def test_whitelist_is_nonempty() -> None:
    assert "sub-fund" in _HYPHEN_WHITELIST
    assert "non-listed" in _HYPHEN_WHITELIST
    assert len(_HYPHEN_WHITELIST) >= 20
