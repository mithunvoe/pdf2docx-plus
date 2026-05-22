"""Tests for the post-emit ``collapse_empty_sections`` pass.

Upstream occasionally emits ``<w:sectPr>`` boundaries between empty
placeholder paragraphs (header-detection stubs, decorative breaks).
Each orphan section forces a page break, so the reader sees a blank
page for every stub. The pass under test removes sections whose body
has no visible content, merging them into the next section.
"""

from __future__ import annotations

import importlib.util
import io
import struct
import zlib
from pathlib import Path

import pytest
from docx import Document  # type: ignore
from docx.opc.constants import RELATIONSHIP_TYPE as RT  # type: ignore
from docx.oxml import OxmlElement  # type: ignore
from docx.oxml.ns import qn  # type: ignore

_SPEC = importlib.util.spec_from_file_location(
    "_sections_under_test",
    Path(__file__).resolve().parent.parent / "pdf2docx_plus" / "emit" / "sections.py",
)
assert _SPEC and _SPEC.loader
_MOD = importlib.util.module_from_spec(_SPEC)
_SPEC.loader.exec_module(_MOD)
collapse_empty_sections = _MOD.collapse_empty_sections


def _append_sect_break(doc) -> None:
    """Insert an empty paragraph that carries a ``<w:sectPr>`` stub,
    placed before the final body-level ``<w:sectPr>``."""
    body = doc.element.body
    final_sect = body.find(qn("w:sectPr"))
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    sectPr = OxmlElement("w:sectPr")
    pPr.append(sectPr)
    p.append(pPr)
    if final_sect is not None:
        final_sect.addprevious(p)
    else:
        body.append(p)


def _section_count(doc) -> int:
    body = doc.element.body
    return sum(1 for _ in body.iter(qn("w:sectPr")))


@pytest.mark.unit
def test_collapses_empty_leading_section() -> None:
    """A section containing only an empty placeholder paragraph is
    removed; the following section takes its place."""
    doc = Document()
    _append_sect_break(doc)  # empty section 1
    doc.add_paragraph("real content")
    before = _section_count(doc)
    collapsed = collapse_empty_sections(doc)
    assert collapsed == 1
    assert _section_count(doc) == before - 1
    assert any(p.text == "real content" for p in doc.paragraphs)


@pytest.mark.unit
def test_collapses_multiple_consecutive_empty_sections() -> None:
    doc = Document()
    for _ in range(3):
        _append_sect_break(doc)
    doc.add_paragraph("body A")
    _append_sect_break(doc)
    doc.add_paragraph("body B")
    collapsed = collapse_empty_sections(doc)
    assert collapsed == 3
    texts = [p.text for p in doc.paragraphs]
    assert "body A" in texts
    assert "body B" in texts


@pytest.mark.unit
def test_preserves_section_with_text() -> None:
    """A section containing a real paragraph must not be collapsed."""
    doc = Document()
    doc.add_paragraph("section-1 content")
    _append_sect_break(doc)
    doc.add_paragraph("section-2 content")
    collapsed = collapse_empty_sections(doc)
    assert collapsed == 0


@pytest.mark.unit
def test_preserves_section_containing_table() -> None:
    """A section whose only visible content is a table must not be collapsed."""
    doc = Document()
    doc.add_table(rows=1, cols=1).cell(0, 0).text = "cell text"
    _append_sect_break(doc)
    doc.add_paragraph("after")
    collapsed = collapse_empty_sections(doc)
    assert collapsed == 0


@pytest.mark.unit
def test_never_removes_final_section() -> None:
    """The final section uses the body-level ``sectPr`` and must be
    preserved even when its content is whitespace only."""
    doc = Document()
    # default new Document has one empty paragraph + body-level sectPr
    collapsed = collapse_empty_sections(doc)
    assert collapsed == 0
    assert _section_count(doc) >= 1


@pytest.mark.unit
def test_idempotent() -> None:
    doc = Document()
    for _ in range(2):
        _append_sect_break(doc)
    doc.add_paragraph("only real content")
    first = collapse_empty_sections(doc)
    second = collapse_empty_sections(doc)
    assert first == 2
    assert second == 0


def _append_drawing_only_section(doc) -> None:
    """Insert a section containing ONLY a single drawing element followed
    by a section break. This mimics the per-page logo image that
    upstream emits before the section's real body content is merged
    elsewhere by stitching."""
    body = doc.element.body
    final_sect = body.find(qn("w:sectPr"))

    # paragraph that hosts the drawing
    p_draw = OxmlElement("w:p")
    r = OxmlElement("w:r")
    drawing = OxmlElement("w:drawing")
    r.append(drawing)
    p_draw.append(r)

    # paragraph that carries the section break
    p_sect = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    sectPr = OxmlElement("w:sectPr")
    pPr.append(sectPr)
    p_sect.append(pPr)

    if final_sect is not None:
        final_sect.addprevious(p_draw)
        final_sect.addprevious(p_sect)
    else:
        body.append(p_draw)
        body.append(p_sect)


@pytest.mark.unit
def test_collapses_drawing_only_section() -> None:
    """A section whose only content is a single decorative drawing
    (typically the per-page header logo) is collapsed."""
    doc = Document()
    _append_drawing_only_section(doc)
    doc.add_paragraph("real content")
    collapsed = collapse_empty_sections(doc)
    assert collapsed == 1
    assert any(p.text == "real content" for p in doc.paragraphs)


@pytest.mark.unit
def test_preserves_section_with_drawing_plus_text() -> None:
    """A section containing both a drawing AND meaningful text is NOT
    collapsed - the logo + text combination is real content."""
    doc = Document()
    body = doc.element.body
    final_sect = body.find(qn("w:sectPr"))
    # paragraph with drawing
    p_draw = OxmlElement("w:p")
    r = OxmlElement("w:r")
    drawing = OxmlElement("w:drawing")
    r.append(drawing)
    p_draw.append(r)
    # paragraph with text + sectPr
    p_text = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    sectPr = OxmlElement("w:sectPr")
    pPr.append(sectPr)
    p_text.append(pPr)
    r2 = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Section text"
    r2.append(t)
    p_text.append(r2)
    if final_sect is not None:
        final_sect.addprevious(p_draw)
        final_sect.addprevious(p_text)
    doc.add_paragraph("trailing")
    collapsed = collapse_empty_sections(doc)
    assert collapsed == 0


@pytest.mark.unit
def test_preserves_section_with_multiple_drawings() -> None:
    """A section with two or more drawings is genuine visual content
    (e.g. a slide-deck-style image gallery) and must be preserved."""
    doc = Document()
    body = doc.element.body
    final_sect = body.find(qn("w:sectPr"))
    # paragraph with two drawings inside two runs
    p_draw = OxmlElement("w:p")
    for _ in range(2):
        r = OxmlElement("w:r")
        drawing = OxmlElement("w:drawing")
        r.append(drawing)
        p_draw.append(r)
    # section-break paragraph
    p_sect = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    sectPr = OxmlElement("w:sectPr")
    pPr.append(sectPr)
    p_sect.append(pPr)
    if final_sect is not None:
        final_sect.addprevious(p_draw)
        final_sect.addprevious(p_sect)
    doc.add_paragraph("body")
    collapsed = collapse_empty_sections(doc)
    assert collapsed == 0


def _png_1x1() -> bytes:
    """A minimal valid 1x1 RGB PNG so python-docx accepts the image."""

    def chunk(typ: bytes, data: bytes) -> bytes:
        body = typ + data
        return (
            struct.pack(">I", len(data)) + body + struct.pack(">I", zlib.crc32(body) & 0xFFFFFFFF)
        )

    sig = b"\x89PNG\r\n\x1a\n"
    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00\xff\x00\x00")
    return sig + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")


def _split_off_drawing_section(doc) -> None:
    """Turn the document's trailing picture paragraph into its own section
    by inserting a sectPr-break paragraph after it, then add a real body
    section so there is something to (potentially) collapse into."""
    body = doc.element.body
    final_sect = body.find(qn("w:sectPr"))
    p_sect = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pPr.append(OxmlElement("w:sectPr"))
    p_sect.append(pPr)
    final_sect.addprevious(p_sect)
    doc.add_paragraph("real content")


@pytest.mark.unit
def test_keeps_drawing_only_section_when_image_is_sole_copy() -> None:
    """A lone drawing that is the *only* copy of a real image must be
    kept - collapsing it would delete the picture entirely (the cover-logo
    regression on KFS-style documents)."""
    doc = Document()
    doc.add_picture(io.BytesIO(_png_1x1()))  # picture lands in its own paragraph
    _split_off_drawing_section(doc)
    before = len(list(doc.element.body.iter(qn("a:blip"))))
    collapsed = collapse_empty_sections(doc)
    after = len(list(doc.element.body.iter(qn("a:blip"))))
    assert before == 1
    assert collapsed == 0
    assert after == 1  # the image survives in the body


@pytest.mark.unit
def test_collapses_drawing_only_section_when_image_lives_in_header() -> None:
    """A lone drawing whose image also appears in a header is a redundant
    per-page logo copy and is collapsed - the header keeps the picture."""
    doc = Document()
    doc.add_picture(io.BytesIO(_png_1x1()))
    body = doc.element.body
    rid = next(iter(body.iter(qn("a:blip")))).get(qn("r:embed"))
    image_part = doc.part.rels[rid].target_part

    # relate the SAME image part into the section header and reference it
    header = doc.sections[0].header
    header.is_linked_to_previous = False
    hrid = header.part.relate_to(image_part, RT.IMAGE)
    run = header.paragraphs[0].add_run()
    drawing = OxmlElement("w:drawing")
    blip = OxmlElement("a:blip")
    blip.set(qn("r:embed"), hrid)
    drawing.append(blip)
    run._r.append(drawing)

    _split_off_drawing_section(doc)
    collapsed = collapse_empty_sections(doc)
    assert collapsed == 1
    assert len(list(doc.element.body.iter(qn("a:blip")))) == 0
