"""Tests for empty-paragraph collapse."""

from __future__ import annotations

import pytest
from docx import Document  # type: ignore

from pdf2docx_plus.emit.whitespace import collapse_empty_paragraphs


@pytest.mark.unit
def test_collapses_consecutive_empties() -> None:
    doc = Document()
    doc.add_paragraph("first")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("second")
    removed = collapse_empty_paragraphs(doc, max_consecutive=1)
    assert removed == 2  # 3 empties -> 1 kept, 2 removed
    texts = [p.text for p in doc.paragraphs]
    assert texts.count("first") == 1
    assert texts.count("second") == 1


@pytest.mark.unit
def test_preserves_paragraph_with_drawing() -> None:
    # add an image-like paragraph via raw XML
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    doc.add_paragraph("a")
    p = doc.add_paragraph()
    r = OxmlElement("w:r")
    d = OxmlElement("w:drawing")
    r.append(d)
    p._p.append(r)
    doc.add_paragraph("b")
    collapse_empty_paragraphs(doc)
    # drawing paragraph must still be present
    drawings = doc.element.body.findall(f".//{qn('w:drawing')}")
    assert len(drawings) == 1


@pytest.mark.unit
def test_empty_body_safe() -> None:
    doc = Document()
    assert collapse_empty_paragraphs(doc) == 0


@pytest.mark.unit
def test_all_empties_collapse_to_max() -> None:
    doc = Document()
    starting = len(doc.paragraphs)  # python-docx Document starts empty but may vary
    for _ in range(5):
        doc.add_paragraph()
    total_empty = starting + 5
    removed = collapse_empty_paragraphs(doc, max_consecutive=1)
    assert removed == total_empty - 1
