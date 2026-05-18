"""Issue P-1: empty-checkbox glyph canonicalisation.

Different Symbol-mapped fonts in the same PDF render visually-empty
checkboxes with different Unicode codepoints — U+2610, U+25A1, U+25FB,
U+25A3, etc.  A redline diff between OLD and NEW PDFs sees character-
level differences for what is semantically the same "empty checkbox".

After P-1, the post-emit pass rewrites all empty-checkbox variants to
the canonical U+25A1 (``□``).  Checked / crossed / filled variants are
left alone because they carry distinct user-intent semantics.
"""

from __future__ import annotations

import importlib.util
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

_SPEC = importlib.util.spec_from_file_location(
    "_checkbox_under_test",
    Path(__file__).resolve().parent.parent
    / "pdf2docx_plus"
    / "emit"
    / "checkbox_glyphs.py",
)
assert _SPEC and _SPEC.loader
_MOD = importlib.util.module_from_spec(_SPEC)
_SPEC.loader.exec_module(_MOD)
canonicalise_checkbox_glyphs = _MOD.canonicalise_checkbox_glyphs
CANONICAL = _MOD.CANONICAL_EMPTY_CHECKBOX  # "□"


def _doc_with_text(text: str) -> Document:
    doc = Document()
    doc.add_paragraph(text)
    return doc


@pytest.mark.unit
@pytest.mark.parametrize(
    "variant",
    [
        "☐",  # BALLOT BOX
        "◻",  # WHITE MEDIUM SQUARE
        "◽",  # WHITE MEDIUM SMALL SQUARE
        "▣",  # WHITE SQUARE CONTAINING BLACK SMALL SQUARE (▣)
        "▫",  # WHITE SMALL SQUARE (▫)
    ],
)
def test_variants_canonicalise_to_canonical_empty(variant: str) -> None:
    doc = _doc_with_text(f"Tick: {variant}")
    changed = canonicalise_checkbox_glyphs(doc)
    assert changed == 1
    assert doc.paragraphs[0].text == f"Tick: {CANONICAL}"


@pytest.mark.unit
def test_canonical_codepoint_is_idempotent() -> None:
    doc = _doc_with_text(f"Tick: {CANONICAL}")
    changed = canonicalise_checkbox_glyphs(doc)
    assert changed == 0
    assert doc.paragraphs[0].text == f"Tick: {CANONICAL}"


@pytest.mark.unit
def test_checked_box_is_preserved() -> None:
    """U+2611 BALLOT BOX WITH CHECK carries user-intent semantics.
    It must NOT collapse to the empty form."""
    doc = _doc_with_text("Tick: ☑")
    changed = canonicalise_checkbox_glyphs(doc)
    assert changed == 0
    assert doc.paragraphs[0].text == "Tick: ☑"


@pytest.mark.unit
def test_crossed_box_is_preserved() -> None:
    """U+2612 BALLOT BOX WITH X carries user-intent semantics."""
    doc = _doc_with_text("Tick: ☒")
    changed = canonicalise_checkbox_glyphs(doc)
    assert changed == 0
    assert doc.paragraphs[0].text == "Tick: ☒"


@pytest.mark.unit
def test_black_square_is_preserved() -> None:
    """U+25A0 BLACK SQUARE is a filled box — distinct from an empty box."""
    doc = _doc_with_text("Tick: ■")
    changed = canonicalise_checkbox_glyphs(doc)
    assert changed == 0
    assert doc.paragraphs[0].text == "Tick: ■"


@pytest.mark.unit
def test_multiple_variants_in_one_run_each_count_once() -> None:
    """Each <w:t> element with at least one rewrite counts as one."""
    doc = _doc_with_text("☐ ◻ ▣ □")
    changed = canonicalise_checkbox_glyphs(doc)
    # Only one w:t (text run) is rewritten; the count is "elements
    # modified", not "characters substituted".
    assert changed == 1
    assert doc.paragraphs[0].text == f"{CANONICAL} {CANONICAL} {CANONICAL} {CANONICAL}"


@pytest.mark.unit
def test_unrelated_text_is_unchanged() -> None:
    doc = _doc_with_text("ordinary text with no checkbox glyphs")
    changed = canonicalise_checkbox_glyphs(doc)
    assert changed == 0


@pytest.mark.unit
def test_empty_document_safe() -> None:
    doc = Document()
    assert canonicalise_checkbox_glyphs(doc) == 0


@pytest.mark.unit
def test_rewrites_text_in_table_cells() -> None:
    doc = Document()
    tbl = doc.add_table(rows=1, cols=2)
    tbl.cell(0, 0).text = "Tick: ☐"
    tbl.cell(0, 1).text = "Tick: ◻"
    changed = canonicalise_checkbox_glyphs(doc)
    # Each cell's paragraph has at least one rewritten <w:t> element.
    assert changed >= 2
    assert tbl.cell(0, 0).text == f"Tick: {CANONICAL}"
    assert tbl.cell(0, 1).text == f"Tick: {CANONICAL}"
