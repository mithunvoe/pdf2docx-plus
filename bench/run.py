"""Benchmark runner.

Usage:
    python -m bench.run --corpus bench/corpus --out bench/reports/latest.json
    python -m bench.run --corpus bench/corpus --baseline bench/reports/v0.5.12.json \\
                       --fail-on-regression 2.0

Corpus layout:
    bench/corpus/<docname>/
        input.pdf
        expected_text.txt      # optional ground-truth text for F1
        expected_tables.json   # optional list[table] for TEDS
        expected_order.json    # optional list[int] for Kendall-tau

Report schema (JSON):
    {
        "pdf2docx_plus_version": "...",
        "generated_at": "...",
        "results": [
            {
                "name": "...",
                "pages": 9,
                "pages_ok": 9,
                "pages_failed": 0,
                "elapsed_s": 4.01,
                "output_bytes": 87084,
                "text_f1": 0.978,
                "teds": null,
                "kendall_tau": null,
                "render_ssim": null,
                "editability": 0.12,
                "warnings": []
            }
        ]
    }
"""

from __future__ import annotations

import argparse
import datetime as _dt
import json
import sys
from pathlib import Path
from typing import Any

from docx import Document  # type: ignore

from bench.metrics import (
    editability,
    kendall_tau,
    render_ssim,
    teds,
    text_char_accuracy,
    text_f1,
)
from pdf2docx_plus import __version__, convert


def _extract_docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _load_json(path: Path) -> Any:
    if not path.exists():
        return None
    with path.open(encoding="utf-8") as f:
        return json.load(f)


def _eval_doc(doc_dir: Path, out_dir: Path, *, compute_ssim: bool) -> dict[str, Any]:
    pdf = doc_dir / "input.pdf"
    out_docx = out_dir / f"{doc_dir.name}.docx"
    out_docx.parent.mkdir(parents=True, exist_ok=True)

    result = convert(pdf, out_docx, timeout_s=240, continue_on_error=True)
    produced_text = _extract_docx_text(out_docx) if out_docx.exists() else ""

    expected_text = None
    exp_txt = doc_dir / "expected_text.txt"
    if exp_txt.exists():
        expected_text = exp_txt.read_text(encoding="utf-8", errors="replace")

    exp_tables = _load_json(doc_dir / "expected_tables.json")
    exp_order = _load_json(doc_dir / "expected_order.json")

    metrics: dict[str, Any] = {
        "name": doc_dir.name,
        "pages": result.pages_total,
        "pages_ok": result.pages_ok,
        "pages_failed": result.pages_failed,
        "elapsed_s": result.elapsed_s,
        "pages_per_second": result.pages_per_second,
        "peak_rss_mb": result.peak_rss_mb,
        "output_bytes": out_docx.stat().st_size if out_docx.exists() else 0,
        "text_f1": text_f1(produced_text, expected_text) if expected_text is not None else None,
        "text_char_accuracy": (
            text_char_accuracy(produced_text, expected_text) if expected_text is not None else None
        ),
        "teds": None,
        "kendall_tau": (
            kendall_tau(exp_order or [], exp_order or []) if exp_order is not None else None
        ),
        "render_ssim": (render_ssim(out_docx, pdf) if compute_ssim and out_docx.exists() else None),
        "editability": editability(out_docx) if out_docx.exists() else None,
        "runs_merged": result.runs_merged,
        "lists_detected": result.lists_detected,
        "headers_footers_detected": result.headers_footers_detected,
        "stitched_table_pairs": len(result.stitched_table_pairs),
        "demoted_floating_images": result.demoted_floating_images,
        "scanned_pages": result.scanned_pages,
        "warnings": list(result.warnings),
    }
    if exp_tables is not None and exp_tables:
        # TEDS only makes sense when ground truth is provided; average across tables
        try:
            doc = Document(str(out_docx))
            produced_tables = [[[c.text for c in row.cells] for row in t.rows] for t in doc.tables]
            n = min(len(produced_tables), len(exp_tables))
            if n:
                scores = [teds(produced_tables[i], exp_tables[i]) for i in range(n)]
                metrics["teds"] = sum(scores) / len(scores)
        except Exception as e:
            metrics["warnings"].append(f"teds failed: {e}")
    return metrics


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(prog="bench.run")
    parser.add_argument("--corpus", type=Path, required=True)
    parser.add_argument("--out", type=Path, default=Path("bench/reports/latest.json"))
    parser.add_argument("--baseline", type=Path, default=None)
    parser.add_argument(
        "--fail-on-regression",
        type=float,
        default=None,
        help="Percent drop in mean text_f1 that fails the run.",
    )
    parser.add_argument("--no-ssim", action="store_true")
    args = parser.parse_args(argv)

    if not args.corpus.is_dir():
        print(f"corpus dir not found: {args.corpus}", file=sys.stderr)
        return 2

    doc_dirs = sorted(
        [p for p in args.corpus.iterdir() if p.is_dir() and (p / "input.pdf").exists()]
    )
    if not doc_dirs:
        print(f"no docs in corpus {args.corpus}", file=sys.stderr)
        return 2

    out_dir = args.out.parent / "outputs"
    results = [_eval_doc(d, out_dir, compute_ssim=not args.no_ssim) for d in doc_dirs]

    report = {
        "pdf2docx_plus_version": __version__,
        "generated_at": _dt.datetime.now(_dt.UTC).isoformat(),
        "results": results,
    }
    args.out.parent.mkdir(parents=True, exist_ok=True)
    args.out.write_text(json.dumps(report, indent=2), encoding="utf-8")

    _print_summary(report)

    if args.fail_on_regression is not None and args.baseline:
        baseline = json.loads(args.baseline.read_text(encoding="utf-8"))
        regressed = _check_regression(report, baseline, args.fail_on_regression)
        if regressed:
            print(f"REGRESSION: {regressed}", file=sys.stderr)
            return 1
    return 0


def _print_summary(report: dict[str, Any]) -> None:
    print(f"\n== pdf2docx-plus {report['pdf2docx_plus_version']} ==")
    print(
        f"{'name':<28} {'pg':>3} {'ok':>3} {'fail':>4} "
        f"{'sec':>6} {'pg/s':>5} {'rss':>5} "
        f"{'f1':>5} {'edit':>5} "
        f"{'runs':>5} {'lists':>5} {'hf':>3} {'stch':>4}"
    )
    print("-" * 96)
    for r in report["results"]:
        f1 = r.get("text_f1")
        edit = r.get("editability")
        rss = r.get("peak_rss_mb") or 0
        print(
            f"{r['name'][:28]:<28} "
            f"{r['pages']:>3} {r['pages_ok']:>3} {r['pages_failed']:>4} "
            f"{r['elapsed_s']:>6.2f} {r['pages_per_second']:>5.2f} "
            f"{int(rss):>5} "
            f"{'  n/a' if f1 is None else f'{f1:>5.3f}'} "
            f"{'  n/a' if edit is None else f'{edit:>5.3f}'} "
            f"{r['runs_merged']:>5} {r['lists_detected']:>5} "
            f"{r['headers_footers_detected']:>3} {r['stitched_table_pairs']:>4}"
        )
    total = sum(r["pages"] for r in report["results"])
    failed = sum(r["pages_failed"] for r in report["results"])
    elapsed = sum(r["elapsed_s"] for r in report["results"])
    total_runs = sum(r["runs_merged"] for r in report["results"])
    total_lists = sum(r["lists_detected"] for r in report["results"])
    print(
        f"\nTOTAL pages={total} failed={failed} elapsed={elapsed:.2f}s "
        f"pg/s={total / elapsed:.2f} runs_merged={total_runs} lists={total_lists}"
    )


def _check_regression(new: dict[str, Any], base: dict[str, Any], pct: float) -> str | None:
    def mean_f1(r: dict[str, Any]) -> float:
        xs = [x["text_f1"] for x in r["results"] if x.get("text_f1") is not None]
        return sum(xs) / len(xs) if xs else 0.0

    if mean_f1(new) + pct / 100 < mean_f1(base):
        return f"mean text_f1 dropped from {mean_f1(base):.3f} to {mean_f1(new):.3f}"
    return None


if __name__ == "__main__":
    sys.exit(main())
