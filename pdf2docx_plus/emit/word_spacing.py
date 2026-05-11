"""Repair word-glue where adjacent runs concatenate without whitespace.

Upstream ``pdf2docx`` rebuilds paragraph text by concatenating the
``TextSpan`` objects that the layout analyser found on each PDF line.
When a sentence wraps across lines in the source PDF, the trailing
space of the preceding line is represented by the line-break itself,
not by a space character inside the span.  Upstream's emitter drops
the line break and emits the two spans as adjacent ``<w:r>`` runs,
producing outputs like ``"confirms,having"`` and ``"Sub-Fund.The"``
where the source clearly read ``"confirms, having"`` and
``"Sub-Fund. The"``.

The post-emit pass here scans every paragraph, looks at adjacent
``<w:r>`` siblings, and inserts a single space between them when the
left run ends with a punctuation character that forces a word break
in the source (comma, semicolon, colon, question mark, exclamation,
closing paren) or with a word-ending period, and the right run
starts with a letter.

What we deliberately leave untouched:

  * Single-letter initials such as ``U.S.`` / ``e.g.`` / ``i.e.`` -
    the left-context regex requires at least two lowercase letters
    before the period, OR a lowercase-then-capital pattern such as
    ``"Sub-Fund."``.
  * Mid-word hyphens.  When the left run ends with ``"-"`` (or a
    hyphen-containing tail like ``"Sub-"``) we never insert a space,
    so a soft-wrap that broke ``"Sub-Fund"`` across two PDF lines is
    rejoined as ``"Sub-Fund"``, not ``"Sub Fund"`` or ``"SubFund"``.
    A small whitelist of commonly-broken financial / legal compound
    words is also enforced so an aggressive future regex can't
    over-fire on them.
"""

from __future__ import annotations

import re
from typing import Any

from docx.oxml.ns import qn  # type: ignore

# Punctuation that unambiguously closes a word when followed by a letter.
_HARD_BREAK = frozenset(",;:?!)")

# Word-ending period regex. Requires at least two lowercase letters
# before the period, OR a lowercase-then-capital pattern, OR a
# hyphenated tail.  Single-letter initials (``U.``, ``S.``) are
# rejected.
_WORD_END_PERIOD = re.compile(r"[a-z]{2}\.$|[a-z][A-Z][a-z]+\.$|[a-z]-[A-Za-z]+\.$")

# Right-context: any run starting with a letter is a candidate; the
# left-context test below guarantees we do not break initials.
_RIGHT_STARTS_LETTER = re.compile(r"^[A-Za-z]")

# Hyphen-protection whitelist.  Specifically protects against
# erroneous de-hyphenation when a hyphenated compound falls at a soft
# line break.  The match is case-insensitive and considers the whole
# hyphenated token, so ``"Sub-Fund"``, ``"sub-fund"``, ``"Sub-Funds"``
# are all guarded.  Add patterns sparingly; over-broad entries will
# mask genuine line-break issues.
_HYPHEN_WHITELIST: frozenset[str] = frozenset({
    "sub-fund",
    "sub-funds",
    "sub-investment",
    "non-deliverable",
    "non-listed",
    "non-residents",
    "non-resident",
    "non-recurring",
    "non-cash",
    "open-ended",
    "closed-ended",
    "co-investor",
    "co-investors",
    "follow-on",
    "follow-up",
    "high-yield",
    "long-term",
    "short-term",
    "near-term",
    "all-in",
    "off-balance",
    "on-balance",
    "over-the-counter",
    "pre-emptive",
    "pre-payment",
    "self-managed",
    "single-strategy",
    "multi-strategy",
    "multi-asset",
    "third-party",
    "year-end",
})


def repair_wrap_spacing(document: Any) -> int:
    """Insert missing spaces between adjacent runs that were joined
    across a PDF soft line break.

    Returns the number of insertions performed.
    """
    fixed = 0
    for paragraph in document.paragraphs:
        fixed += _repair_paragraph(paragraph._p)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    fixed += _repair_paragraph(paragraph._p)
    return fixed


def _repair_paragraph(p_elem: Any) -> int:
    runs = p_elem.findall(qn("w:r"))
    if len(runs) < 2:
        return 0
    fixed = 0
    for i in range(len(runs) - 1):
        a, b = runs[i], runs[i + 1]
        if a.getnext() is not b:
            # something between the two runs (hyperlink, bookmark, br)
            continue
        if _contains_line_break(a) or _contains_line_break(b):
            continue
        left_text = _last_t_text(a)
        right_text = _first_t_text(b)
        if not left_text or not right_text:
            continue
        if left_text.endswith(" ") or right_text.startswith(" "):
            continue
        if not _RIGHT_STARTS_LETTER.match(right_text):
            continue
        # Never split across a hyphen: ``"Sub-"|"Fund"`` must stay
        # ``"Sub-Fund"``.  Cover both the trailing-hyphen literal and
        # the whitelist of common hyphenated compounds.
        if left_text.endswith("-"):
            continue
        if _is_protected_hyphen_pair(left_text, right_text):
            continue
        last_char = left_text[-1]
        if last_char in _HARD_BREAK:
            _append_space_to_last_t(a)
            fixed += 1
            continue
        if last_char == "." and _WORD_END_PERIOD.search(left_text):
            _append_space_to_last_t(a)
            fixed += 1
            continue
    return fixed


def _is_protected_hyphen_pair(left_text: str, right_text: str) -> bool:
    """Return True when ``<left><right>`` would form a whitelisted
    hyphenated token.

    Both halves are stripped of leading/trailing whitespace, then
    joined and compared case-insensitively against the whitelist.
    """
    if "-" not in left_text and "-" not in right_text:
        return False
    # take the rightmost hyphen-containing word of left and the
    # leftmost hyphen-containing word of right
    left_word = re.split(r"\s+", left_text.strip())[-1] if left_text.strip() else ""
    right_word = re.split(r"\s+", right_text.strip())[0] if right_text.strip() else ""
    if not ("-" in left_word or "-" in right_word):
        return False
    joined = (left_word + right_word).lower().strip()
    if joined in _HYPHEN_WHITELIST:
        return True
    # also accept plural/genitive forms by stripping trailing 's'
    if joined.endswith("s") and joined[:-1] in _HYPHEN_WHITELIST:
        return True
    return False


def _contains_line_break(run: Any) -> bool:
    return run.find(qn("w:br")) is not None or run.find(qn("w:tab")) is not None


def _last_t_text(run: Any) -> str:
    ts = run.findall(qn("w:t"))
    if not ts:
        return ""
    t = ts[-1]
    return t.text or ""


def _first_t_text(run: Any) -> str:
    ts = run.findall(qn("w:t"))
    if not ts:
        return ""
    t = ts[0]
    return t.text or ""


def _append_space_to_last_t(run: Any) -> None:
    ts = run.findall(qn("w:t"))
    if not ts:
        return
    t = ts[-1]
    t.text = (t.text or "") + " "
    t.set(qn("xml:space"), "preserve")
