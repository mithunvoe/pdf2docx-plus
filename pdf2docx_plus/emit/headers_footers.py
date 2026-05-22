"""Move paragraphs detected as headers/footers into section hdr/ftr parts.

Upstream ``pdf2docx`` writes the repeating top-of-page and bottom-of-page
text as ordinary body paragraphs.  That is wrong twice over:

  * each repetition renders inline in the body, so a 60-page document
    shows the same banner 60 times in the flow;
  * the *real* page header/footer parts of every section are left
    empty, so when Word repaginates the original chrome disappears.

This pass walks ``doc.paragraphs`` and, for each
``HeaderFooter`` candidate detected by ``layout.hf_detect``:

  1. identifies the set of ``w:p`` elements in the body whose
     normalised text matches the candidate;
  2. writes ONE representative copy into the appropriate
     ``section.header`` / ``section.footer`` of the section the page
     range belongs to (multi-section aware - we no longer blanket
     section[0]);
  3. removes the remaining matching body paragraphs.

If the document has only one section we behave exactly like the
previous implementation - a single representative paragraph lifts into
``section[0].header`` / ``section[0].footer``.  When upstream emits the
typical "one section per source PDF page" layout, we group consecutive
sections that share the same detected text into ranges, link the later
ones to the previous via ``is_linked_to_previous`` and only write the
content once.
"""

from __future__ import annotations

import re
from collections import defaultdict
from typing import Any

from docx.oxml.ns import qn  # type: ignore

from ..layout.hf_detect import HeaderFooter

_PAGE_NUM = re.compile(r"(?:page\s*)?\d+(?:\s*/\s*\d+)?", re.IGNORECASE)
_WS = re.compile(r"\s+")
# Footer lines that embed a page number ("N Last update: <date>") are
# deferred to ``promote_page_numbers_to_footer`` so the page number
# becomes a live ``PAGE`` field instead of a static digit copied onto
# every page. Extracting them here would freeze "1" on all pages.
_DEFERRED_TO_PAGE_FOOTER = re.compile(r"last\s*update", re.IGNORECASE)


def _norm(text: str) -> str:
    return _WS.sub(" ", _PAGE_NUM.sub("#", text)).strip()


def _is_deferred_footer(text: str) -> bool:
    return bool(_DEFERRED_TO_PAGE_FOOTER.search(text))


def extract_headers_footers(doc: Any, detected: list[HeaderFooter]) -> int:
    """Move matching body paragraphs into per-section hdr/ftr parts.

    Returns the number of body paragraphs removed.

    When the document has one section, the behaviour is unchanged from
    the previous implementation: one canonical paragraph lands in
    ``section[0].header`` / ``section[0].footer``, every other matching
    body paragraph is deleted.

    When the document has multiple sections (upstream emits one section
    per source PDF page by default), each section's first body
    paragraph that matches a header/footer triggers a write into that
    section's header/footer; subsequent sections in the document order
    that share the same content are linked-to-previous so the rendered
    DOCX only stores one copy of the chrome.
    """
    if not detected:
        return 0

    headers = [h for h in detected if h.is_header and _is_meaningful(h.text)]
    footers = [
        h
        for h in detected
        if not h.is_header
        and _is_meaningful(h.text)
        and not _is_deferred_footer(h.text)
    ]
    header_texts = {_norm(h.text) for h in headers}
    footer_texts = {_norm(h.text) for h in footers}
    if not header_texts and not footer_texts:
        return 0

    sections = list(doc.sections)
    if not sections:
        return 0

    # Build per-section paragraph buckets. The body is a sequence of
    # paragraphs and tables; a paragraph that carries ``<w:sectPr>``
    # is the LAST paragraph of its section.
    body = doc.element.body
    section_buckets: list[list[Any]] = [[]]
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            section_buckets[-1].append(child)
            if child.find(qn("w:pPr") + "/" + qn("w:sectPr")) is not None:
                section_buckets.append([])
    # last (open) bucket belongs to the final section that owns body-level sectPr
    if section_buckets and not section_buckets[-1]:
        section_buckets.pop()
    # align bucket count to section count (defensively)
    while len(section_buckets) < len(sections):
        section_buckets.append([])

    moved = 0
    # Per-section: keep a representative paragraph and remove duplicates.
    # Track which header/footer text we have already installed into each
    # section to avoid double-inserting into the same section.
    written_header: dict[int, set[str]] = defaultdict(set)
    written_footer: dict[int, set[str]] = defaultdict(set)
    # Track the last-seen content per kind so consecutive matching sections
    # can be linked-to-previous (collapses the chrome into one stored copy).
    last_header_text: str | None = None
    last_footer_text: str | None = None

    for idx, bucket in enumerate(section_buckets):
        if idx >= len(sections):
            break
        section = sections[idx]
        header_targets: list[tuple[Any, str]] = []
        footer_targets: list[tuple[Any, str]] = []
        for p in bucket:
            text = _norm(_paragraph_text(p))
            if not text:
                continue
            if text in header_texts:
                header_targets.append((p, text))
            elif text in footer_texts:
                footer_targets.append((p, text))

        # write at most ONE header paragraph into this section
        if header_targets:
            rep, rep_text = header_targets[0]
            if rep_text not in written_header[idx]:
                if rep_text == last_header_text:
                    # already in the previous section's header - link
                    try:
                        section.header.is_linked_to_previous = True
                    except Exception:
                        _copy_paragraph_into(section.header, rep)
                else:
                    _copy_paragraph_into(section.header, rep)
                    last_header_text = rep_text
                written_header[idx].add(rep_text)
            # remove every header paragraph from the body
            for p, _ in header_targets:
                parent = p.getparent()
                if parent is not None:
                    parent.remove(p)
                    moved += 1
        if footer_targets:
            rep, rep_text = footer_targets[0]
            if rep_text not in written_footer[idx]:
                if rep_text == last_footer_text:
                    try:
                        section.footer.is_linked_to_previous = True
                    except Exception:
                        _copy_paragraph_into(section.footer, rep)
                else:
                    _copy_paragraph_into(section.footer, rep)
                    last_footer_text = rep_text
                written_footer[idx].add(rep_text)
            for p, _ in footer_targets:
                parent = p.getparent()
                if parent is not None:
                    parent.remove(p)
                    moved += 1

    return moved


def _is_meaningful(text: str) -> bool:
    """Filter out HF candidates too short or numeric to safely extract.

    Anything shorter than 8 characters after normalisation is almost
    always a page number or pagination artefact; leaving it in the
    body is cheaper than ghosting it across every page.
    """
    norm = _norm(text)
    if len(norm) < 8:
        return False
    return not all(c in "#/ .-:" for c in norm)


def _paragraph_text(p: Any) -> str:
    return "".join((t.text or "") for t in p.iter(qn("w:t")))


def _copy_paragraph_into(target: Any, p_element: Any) -> None:
    """Append the paragraph's run content into a section.header/footer.

    Re-uses the existing first paragraph slot when it is empty
    (python-docx populates a placeholder ``<w:p/>`` even on a fresh
    header), otherwise adds a new paragraph.
    """
    dst_paragraphs = target.paragraphs
    if dst_paragraphs and not dst_paragraphs[0].text.strip():
        dst = dst_paragraphs[0]
    else:
        dst = target.add_paragraph()
    # Lift runs - copy text and basic formatting without touching the
    # source paragraph (which the caller will remove from the body).
    for r in p_element.findall(qn("w:r")):
        text = "".join((t.text or "") for t in r.findall(qn("w:t")))
        if not text:
            continue
        new_run = dst.add_run(text)
        rPr = r.find(qn("w:rPr"))
        if rPr is None:
            continue
        # carry over the most useful run properties
        b = rPr.find(qn("w:b"))
        if b is not None:
            new_run.bold = True
        i = rPr.find(qn("w:i"))
        if i is not None:
            new_run.italic = True
        u = rPr.find(qn("w:u"))
        if u is not None:
            new_run.underline = True
        sz = rPr.find(qn("w:sz"))
        if sz is not None:
            try:
                from docx.shared import Pt  # type: ignore

                new_run.font.size = Pt(int(sz.get(qn("w:val"))) / 2)
            except Exception:
                pass
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is not None:
            name = rFonts.get(qn("w:ascii")) or rFonts.get(qn("w:hAnsi"))
            if name:
                new_run.font.name = name
