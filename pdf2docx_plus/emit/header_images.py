"""Move repeating per-page letterhead images into the section header.

Upstream ``pdf2docx`` re-emits the page-top letterhead / logo image as
an inline (or, worse, floating ``wp:anchor``) drawing on *every* source
PDF page, left in the document **body**.  In the default profile each
PDF page becomes its own DOCX section, so the same logo appears once per
section in the body flow.  That is wrong in three ways:

  * a floating-anchor logo overlaps the first body line (the logo and
    the text are positioned independently, so they collide);
  * the logo repeats inline dozens of times instead of rendering once
    per page as real page chrome;
  * a section whose body is *only* the logo plus a stray page number
    survives section-collapse and renders as a near-blank page.

This pass detects the repeating letterhead by content (the image bytes
are deduplicated by upstream, but we hash the blob so detection also
works when a copy carries a distinct relationship id), lifts one copy
into each carrying section's ``w:hdr`` part — linking consecutive
identical sections via ``is_linked_to_previous`` so only one copy is
stored — and removes the inline body copies.  Emptied chrome sections
then collapse downstream, eliminating the phantom blank pages.

A figure that appears on a single section (a one-off chart) is left in
the body: only an image whose blob recurs across at least
``min(max(2, ratio * n_sections))`` sections is treated as chrome.
"""

from __future__ import annotations

import contextlib
import hashlib
from collections import defaultdict
from io import BytesIO
from typing import Any

from docx.oxml.ns import qn  # type: ignore
from docx.shared import Emu, Pt  # type: ignore

# A letterhead placed in the header is scaled so it never exceeds this
# height; otherwise a tall logo (the source bbox often includes large
# transparent margins) makes the renderer grow the header band and
# inflate the page count. ~1.25in comfortably fits a banner logo.
_MAX_HEADER_IMG_EMU = int(Pt(90))
# gap reserved between the header logo and the first body line.
_HEADER_BODY_GAP_EMU = int(Pt(10))
_DEFAULT_HEADER_DIST_EMU = int(Pt(36))


def promote_header_images_to_section(
    document: Any,
    *,
    ratio: float = 0.1,
    min_sections: int = 2,
    max_height_frac: float = 0.5,
) -> int:
    """Lift the repeating letterhead image into each section's header.

    Args:
        ratio: minimum fraction of sections a top lone image must appear
            on to be treated as a letterhead.
        min_sections: absolute floor for the repetition count, so short
            documents (2-3 pages) still qualify.
        max_height_frac: skip images taller than this fraction of the
            page height — a full-page background or watermark is not a
            header and would overlap the body if placed in the header.

    Only a lone image that is the **first visible block** of its section
    is considered (a real letterhead sits at the top of the page); a
    figure repeated mid-page is left in the body.

    Returns the number of inline body drawings removed.
    """
    sections = list(document.sections)
    if len(sections) < 2:
        return 0

    body = document.element.body
    buckets = _section_buckets(body)
    if len(buckets) < 2:
        return 0
    # align bucket count to section count defensively
    n = min(len(buckets), len(sections))

    page_height_pt = _page_height_pt(sections[0])

    # bucket index -> list of (paragraph, blob_hash, blob, cx, cy, align)
    lone_by_bucket: dict[int, list[tuple]] = defaultdict(list)
    hash_buckets: dict[str, set[int]] = defaultdict(set)
    for idx in range(n):
        seen_text = False
        for p in buckets[idx]:
            if _has_visible_text(p):
                seen_text = True
                continue
            if seen_text:
                # only top-of-section lone images are letterhead chrome
                continue
            info = _lone_image_info(document, p, page_height_pt, max_height_frac)
            if info is None:
                continue
            blob_hash = info[0]
            lone_by_bucket[idx].append((p, *info))
            hash_buckets[blob_hash].add(idx)

    if not hash_buckets:
        return 0

    threshold = max(min_sections, int(len(sections) * ratio + 0.9999))
    letterhead_hashes = {h for h, idxs in hash_buckets.items() if len(idxs) >= threshold}
    if not letterhead_hashes:
        return 0

    removed = 0
    prev_written_hash: str | None = None
    for idx in range(n):
        section = sections[idx]
        carriers = [
            entry for entry in lone_by_bucket.get(idx, []) if entry[1] in letterhead_hashes
        ]
        if not carriers:
            # break the inheritance chain so a non-letterhead page does
            # not inherit the previous section's logo header.
            if prev_written_hash is not None:
                with contextlib.suppress(Exception):
                    section.header.is_linked_to_previous = False
                prev_written_hash = None
            continue

        _p, blob_hash, blob, cx, cy, align = carriers[0]
        if blob_hash == prev_written_hash:
            try:
                section.header.is_linked_to_previous = True
            except Exception:
                _write_image_header(section, blob, cx, cy, align)
        else:
            if _write_image_header(section, blob, cx, cy, align):
                prev_written_hash = blob_hash

        # remove every letterhead lone-image paragraph in this bucket
        for entry in carriers:
            para = entry[0]
            parent = para.getparent()
            if parent is not None:
                parent.remove(para)
                removed += 1

    return removed


# -- detection helpers ----------------------------------------------------


def _section_buckets(body: Any) -> list[list[Any]]:
    """Partition the body's direct ``<w:p>`` children into per-section
    buckets. The paragraph carrying a ``<w:sectPr>`` ends its bucket;
    the final (open) bucket belongs to the body-level sectPr section."""
    buckets: list[list[Any]] = [[]]
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            buckets[-1].append(child)
            if child.find(qn("w:pPr") + "/" + qn("w:sectPr")) is not None:
                buckets.append([])
    if buckets and not buckets[-1]:
        buckets.pop()
    return buckets


def _lone_image_info(
    document: Any, paragraph: Any, page_height_pt: float, max_height_frac: float
) -> tuple[str, bytes, str | None, str | None, str | None] | None:
    """Return ``(blob_hash, blob, cx, cy, jc)`` when the paragraph is a
    lone-image paragraph (no visible text, exactly one drawing) whose
    image qualifies as page chrome; otherwise ``None``."""
    if _has_visible_text(paragraph):
        return None
    drawings = paragraph.findall(".//" + qn("w:drawing"))
    if len(drawings) != 1:
        return None
    drawing = drawings[0]
    blip = drawing.find(".//" + qn("a:blip"))
    if blip is None:
        return None
    embed = blip.get(qn("r:embed"))
    if not embed:
        return None
    blob = _blob_for(document, embed)
    if blob is None:
        return None
    ext = drawing.find(".//" + qn("wp:extent"))
    cx = ext.get("cx") if ext is not None else None
    cy = ext.get("cy") if ext is not None else None
    # skip oversized images that would not behave as a header band
    if cy and page_height_pt > 0:
        try:
            cy_pt = int(cy) / 12700.0
            if cy_pt > max_height_frac * page_height_pt:
                return None
        except (TypeError, ValueError):
            pass
    jc_el = paragraph.find(qn("w:pPr") + "/" + qn("w:jc"))
    jc = jc_el.get(qn("w:val")) if jc_el is not None else None
    return (hashlib.sha1(blob).hexdigest(), blob, cx, cy, jc)


def _has_visible_text(paragraph: Any) -> bool:
    return any((t.text or "").strip() for t in paragraph.iter(qn("w:t")))


def _blob_for(document: Any, embed_id: str) -> bytes | None:
    try:
        part = document.part.related_parts[embed_id]
    except (KeyError, AttributeError):
        return None
    try:
        return part.blob
    except Exception:
        return None


def _page_height_pt(section: Any) -> float:
    try:
        h = section.page_height
        if h is None:
            return 0.0
        return int(h) / 12700.0
    except Exception:
        return 0.0


# -- header writing -------------------------------------------------------


def _write_image_header(
    section: Any, blob: bytes, cx: str | None, cy: str | None, jc: str | None
) -> bool:
    """Place ``blob`` as an inline picture in ``section.header``.

    Uses ``add_picture`` on a header run so the image relationship is
    created in the header part (moving the raw ``<w:drawing>`` element
    would dangle its body-part relationship). Returns True on success.
    """
    header = section.header
    try:
        header.is_linked_to_previous = False
    except Exception:
        return False
    # reuse the placeholder empty paragraph if present
    paras = header.paragraphs
    if (
        paras
        and not paras[0].text.strip()
        and not list(paras[0]._element.iter(qn("w:drawing")))
    ):
        p = paras[0]
    else:
        p = header.add_paragraph()
    run = p.add_run()
    width = height = None
    if cx and cy:
        try:
            cxi, cyi = int(cx), int(cy)
            if cyi > _MAX_HEADER_IMG_EMU and cyi > 0:
                scale = _MAX_HEADER_IMG_EMU / cyi
                cxi = int(cxi * scale)
                cyi = int(cyi * scale)
            width, height = Emu(cxi), Emu(cyi)
        except (TypeError, ValueError):
            width = height = None
    try:
        if width is not None and height is not None:
            run.add_picture(BytesIO(blob), width=width, height=height)
        else:
            run.add_picture(BytesIO(blob))
    except Exception:
        return False
    if jc:
        _set_alignment(p, jc)
    if height is not None:
        _reserve_top_margin(section, int(height))
    return True


def _reserve_top_margin(section: Any, image_height_emu: int) -> None:
    """Grow the section top margin so the body clears the header logo.

    Mirrors the source PDF, where the body text starts below the
    letterhead band. Reserving the space in ``w:top`` keeps the logo and
    the first body line from overlapping without relying on the
    renderer's (inconsistent) header auto-grow behaviour.
    """
    try:
        header_dist = (
            int(section.header_distance)
            if section.header_distance
            else _DEFAULT_HEADER_DIST_EMU
        )
        needed = header_dist + image_height_emu + _HEADER_BODY_GAP_EMU
        current_top = int(section.top_margin) if section.top_margin else 0
        if needed > current_top:
            section.top_margin = Emu(needed)
    except Exception:
        pass


def _set_alignment(paragraph: Any, jc: str) -> None:
    pPr = paragraph._element.get_or_add_pPr()
    existing = pPr.find(qn("w:jc"))
    if existing is None:
        existing = pPr.makeelement(qn("w:jc"), {})
        pPr.append(existing)
    existing.set(qn("w:val"), jc)
