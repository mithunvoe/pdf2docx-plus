"""Image + vector-graph recovery.

Strategy:

* Walk the PDF page-by-page using `fitz`.
* For each page whose xobject-referenced image has no bbox returned by
  PyMuPDF (the "repeated logo" case — upstream silently drops these),
  clip-render the template bbox from a sibling page and inject a new
  paragraph with the image **at the body position corresponding to
  that page**, anchored on upstream's emitted `<w:sectPr>` elements.
* For each page, optionally rasterize clusters of vector drawings
  that do NOT overlap significantly with text blocks (to avoid
  rendering text as a low-DPI raster mess).
"""

from __future__ import annotations

import io
from dataclasses import dataclass, field
from typing import Any

import fitz  # type: ignore
from docx import Document  # type: ignore
from docx.oxml.ns import qn
from docx.shared import Emu


@dataclass
class RecoveryReport:
    missing_raster_recovered: int = 0
    vector_regions_rasterized: int = 0
    pages_touched: list[int] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)


def recover_images(
    pdf_path: str,
    docx_path: str,
    *,
    rasterize_vectors: bool = False,
    recover_missing_rasters: bool = True,
    min_drawing_density: float = 0.05,
    max_text_overlap: float = 0.20,
    render_dpi: int = 150,
) -> RecoveryReport:
    """Patch `docx_path` with recovered images from `pdf_path`.

    Args:
        rasterize_vectors: if True, rasterize dense vector regions.
        recover_missing_rasters: if True, re-emit raster images upstream
            dropped (common for repeated logos).
        min_drawing_density: (drawing_area / cluster_area) ratio below
            which a cluster is ignored.
        max_text_overlap: if a cluster bbox overlaps more than this
            fraction with any text block, skip rasterizing it
            (otherwise we'd render crisp text as a blurry PNG).
        render_dpi: target DPI for rasterization.
    """
    report = RecoveryReport()
    if not (rasterize_vectors or recover_missing_rasters):
        return report

    pdf = fitz.open(pdf_path)
    doc = Document(docx_path)
    try:
        page_anchors = _build_page_anchors(doc, len(pdf))

        for page_index in range(len(pdf)):
            page = pdf[page_index]

            if recover_missing_rasters:
                for img_info in page.get_images(full=True):
                    xref = img_info[0]
                    try:
                        rects = page.get_image_rects(xref)
                    except Exception:
                        rects = []
                    if rects:
                        continue
                    template_bbox = _find_template_bbox(pdf, xref)
                    if template_bbox is None:
                        continue
                    png_bytes = _clip_page_png(page, template_bbox, scale=3.0)
                    if png_bytes is None:
                        continue
                    if _insert_at_page(doc, page_anchors, page_index, png_bytes, template_bbox):
                        report.missing_raster_recovered += 1
                        report.pages_touched.append(page_index)

            if rasterize_vectors:
                text_blocks = _text_block_rects(page)
                for region in _vector_clusters(page, min_drawing_density):
                    if _overlap_frac(region, text_blocks) > max_text_overlap:
                        continue
                    try:
                        pix = page.get_pixmap(
                            clip=region,
                            matrix=fitz.Matrix(render_dpi / 72, render_dpi / 72),
                        )
                        png_bytes = pix.tobytes("png")
                    except Exception as e:
                        report.warnings.append(f"page {page_index + 1} vector raster failed: {e}")
                        continue
                    if _insert_at_page(doc, page_anchors, page_index, png_bytes, region):
                        report.vector_regions_rasterized += 1
                        report.pages_touched.append(page_index)

        if report.missing_raster_recovered or report.vector_regions_rasterized:
            doc.save(docx_path)
    finally:
        pdf.close()
    return report


# -- positioning: map PDF page index -> DOCX body anchor ------------------


def _build_page_anchors(doc: Any, page_count: int) -> list[Any]:
    """Return a list of `page_count` XML elements; the N-th is where
    recovered content for page N should be inserted *before*.

    Upstream emits `<w:sectPr>` between pages. We collect those in body
    order; if we have fewer than `page_count`, we pad with body end.
    """
    body = doc.element.body
    sect_prs: list[Any] = []
    for p in body.iter(qn("w:p")):
        # sectPr lives inside pPr; iter yields descendants in document order
        sp = p.find(f"{qn('w:pPr')}/{qn('w:sectPr')}")
        if sp is not None:
            sect_prs.append(p)
    # also consider top-level sectPr (the final section for the whole doc)
    top_level = body.find(qn("w:sectPr"))
    if top_level is not None:
        sect_prs.append(top_level)

    # anchors: use first `page_count` sectPr-bearing paragraphs.
    # for any page beyond that, anchor at the last known sectPr so content
    # still lands near the end rather than being lost.
    anchors: list[Any] = []
    for i in range(page_count):
        if i < len(sect_prs):
            anchors.append(sect_prs[i])
        elif sect_prs:
            anchors.append(sect_prs[-1])
        else:
            anchors.append(None)  # no anchors at all -> append at end
    return anchors


def _insert_at_page(
    doc: Any,
    anchors: list[Any],
    page_index: int,
    png_bytes: bytes,
    bbox: fitz.Rect,
) -> bool:
    """Create a paragraph with the image and insert it before the page anchor."""
    # build an image paragraph at the end first (python-docx API limitation),
    # then relocate its XML to the correct position.
    p = doc.add_paragraph()
    run = p.add_run()
    width_emu = Emu(int(bbox.width * 9525))
    try:
        run.add_picture(io.BytesIO(png_bytes), width=width_emu)
    except Exception:
        doc.element.body.remove(p._p)
        return False

    anchor = anchors[page_index] if page_index < len(anchors) else None
    if anchor is None:
        return True  # added at end by default
    body = doc.element.body
    # detach from end
    body.remove(p._p)
    # insert before anchor
    anchor.addprevious(p._p)
    return True


# -- template bbox lookup --------------------------------------------------


def _find_template_bbox(pdf: fitz.Document, xref: int) -> fitz.Rect | None:
    for p in pdf:
        try:
            rects = p.get_image_rects(xref)
        except Exception:
            continue
        if rects:
            return rects[0]
    return None


def _clip_page_png(page: fitz.Page, bbox: fitz.Rect, *, scale: float = 3.0) -> bytes | None:
    try:
        pix = page.get_pixmap(clip=bbox, matrix=fitz.Matrix(scale, scale))
        return pix.tobytes("png")
    except Exception:
        return None


# -- vector clustering -----------------------------------------------------


def _text_block_rects(page: fitz.Page) -> list[fitz.Rect]:
    out: list[fitz.Rect] = []
    try:
        raw = page.get_text("dict")
    except Exception:
        return out
    for block in raw.get("blocks", []) or []:
        if block.get("type") != 0:  # 0 = text block
            continue
        b = block.get("bbox")
        if b is None:
            continue
        out.append(fitz.Rect(b))
    return out


def _vector_clusters(page: fitz.Page, min_density: float) -> list[fitz.Rect]:
    """Return bbox rects of clusters of vector drawings likely to be graphics.

    A cluster qualifies only when:

    * its bounding box is at least 100x100 pt (raised from 60x60 -
      avoids rasterizing decorative dividers and small icon clusters),
    * the per-drawing area density inside the cluster is at least
      ``min_density``,
    * the cluster contains a meaningful number of *fill* drawings
      (>= 4) - this filters table-border-only structures whose
      strokes happen to cluster geometrically; a real chart / icon
      has filled regions.

    Cluster bbox is computed from the geometric union of
    individual drawing rects (with a small pad so adjacent shapes
    merge), not from the page geometry.
    """
    drawings = page.get_drawings()
    if not drawings:
        return []
    drawing_meta: list[tuple[fitz.Rect, float, bool]] = []  # rect, area, is_fill
    for d in drawings:
        r = d.get("rect")
        if r is None:
            continue
        rr = fitz.Rect(r)
        if rr.width < 3 or rr.height < 3:
            continue
        d_type = d.get("type", "")
        is_fill = d_type in ("f", "fs")  # fill or fill+stroke
        drawing_meta.append((rr, rr.width * rr.height, is_fill))
    if not drawing_meta:
        return []

    merged = _merge_overlapping([r for r, _, _ in drawing_meta], pad=10.0)

    clusters: list[fitz.Rect] = []
    for m in merged:
        # drawing area + fill count inside this merged bbox
        total_draw_area = 0.0
        fill_count = 0
        for r, a, is_fill in drawing_meta:
            if not (m.contains(r) or m.intersects(r)):
                continue
            total_draw_area += a
            if is_fill:
                fill_count += 1
        m_area = max(m.width * m.height, 1.0)
        density = total_draw_area / m_area
        if m.width < 100 or m.height < 100:
            continue
        if density < min_density:
            continue
        if fill_count < 4:
            # too few filled regions for a chart / diagram
            continue
        clusters.append(m)
    return clusters


def _merge_overlapping(rects: list[fitz.Rect], *, pad: float = 0.0) -> list[fitz.Rect]:
    changed = True
    current = [fitz.Rect(r.x0 - pad, r.y0 - pad, r.x1 + pad, r.y1 + pad) for r in rects]
    while changed:
        changed = False
        out: list[fitz.Rect] = []
        for r in current:
            for i, merged in enumerate(out):
                if merged.intersects(r):
                    out[i] = merged | r
                    changed = True
                    break
            else:
                out.append(fitz.Rect(r))
        current = out
    return current


def _overlap_frac(region: fitz.Rect, blocks: list[fitz.Rect]) -> float:
    """Total fraction of ``region`` area covered by *any* text block.

    Previously this returned the max overlap of a single text block,
    which under-counted dense regions where many small text blocks
    each cover < 20% individually but together cover most of the
    region.  That under-counting is the root cause of "highlighted
    table cells rasterized as a chart": every cell contains a small
    text block that doesn't individually exceed 20% of the cluster
    area, so the cluster was treated as a pure vector graphic.

    The total-area calculation is conservative against double counting:
    overlapping text-block bboxes contribute the union of their
    intersections with ``region``, not the sum.
    """
    region_area = max(region.width * region.height, 1.0)
    inters: list[fitz.Rect] = []
    for b in blocks:
        inter = region & b
        if inter.is_empty:
            continue
        inters.append(inter)
    if not inters:
        return 0.0
    # union area: approximate by summing pixel-level area; for tight
    # bounds we'd run a sweep-line, but a tolerant approximation is
    # enough to filter false positives here.
    union_area = _union_area(inters)
    return min(union_area / region_area, 1.0)


def _union_area(rects: list[fitz.Rect]) -> float:
    """Return an approximate union area of ``rects``.

    Uses a simple inclusion-exclusion truncated at first-order
    (sum of areas minus sum of pairwise intersections), which is exact
    when intersections themselves don't overlap and is a slight
    over-estimate otherwise.  Over-estimating the text coverage is the
    safe direction here: it errs towards NOT rasterizing.
    """
    n = len(rects)
    if n == 0:
        return 0.0
    total = sum(max(r.width, 0) * max(r.height, 0) for r in rects)
    for i in range(n):
        for j in range(i + 1, n):
            inter = rects[i] & rects[j]
            if inter.is_empty:
                continue
            total -= max(inter.width, 0) * max(inter.height, 0)
    return max(total, 0.0)
