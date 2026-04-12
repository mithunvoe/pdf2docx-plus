"""Run consolidation: merge adjacent python-docx runs with identical formatting.

Upstream emits one `<w:r>` per pdf2docx `TextSpan`, which can produce
dozens of runs inside a single paragraph. Word still renders them, but:

* the doc isn't editable (changing a word touches N runs).
* file size bloats.
* the `editability` bench metric stays near zero.

This post-processor walks every paragraph after `make_docx` has finished,
compares adjacent runs' `rPr` XML, and merges them when identical.
"""

from __future__ import annotations

from typing import Any

from docx.oxml.ns import qn  # type: ignore


def consolidate_runs(document: Any) -> int:
    """Merge adjacent runs with identical rPr. Returns number merged."""
    merged = 0
    for paragraph in document.paragraphs:
        merged += _consolidate_paragraph(paragraph._p)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    merged += _consolidate_paragraph(paragraph._p)
    return merged


def _consolidate_paragraph(p_elem: Any) -> int:
    runs = p_elem.findall(qn("w:r"))
    if len(runs) < 2:
        return 0
    merged = 0
    i = 0
    while i < len(runs) - 1:
        a = runs[i]
        b = runs[i + 1]
        # only merge when a and b are adjacent siblings inside the paragraph
        # (hyperlinks / bookmarks between them must break the chain)
        if a.getnext() is not b:
            i += 1
            continue
        if not _rpr_equal(a, b):
            i += 1
            continue
        # append all non-rPr children of b to a
        for child in list(b):
            if child.tag == qn("w:rPr"):
                continue
            a.append(child)
        p_elem.remove(b)
        runs.pop(i + 1)
        merged += 1
        # stay at i to try merging with the new neighbour
    return merged


def _rpr_equal(a: Any, b: Any) -> bool:
    rp_a = a.find(qn("w:rPr"))
    rp_b = b.find(qn("w:rPr"))
    if rp_a is None and rp_b is None:
        return True
    if rp_a is None or rp_b is None:
        return False
    # string-compare canonicalised XML
    from lxml import etree  # type: ignore

    return etree.tostring(rp_a, method="c14n") == etree.tostring(rp_b, method="c14n")
