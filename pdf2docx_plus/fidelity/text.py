"""Text sanitisation: strip XML-1.0 invalid control characters.

The OOXML reader in Word rejects the document outright (or silently
truncates) when the inner `<w:t>` text contains any of the invalid
control characters defined by the XML 1.0 spec: 0x00-0x08, 0x0B, 0x0C,
0x0E-0x1F. Upstream already has an `INVALID_CHARS` string in
`pdf2docx.common.constants` but only applies it in select call sites,
leaving the NULL-byte corruption path open (upstream #324).

We install a paragraph-level `add_run` interceptor that scrubs the text
argument before it is stored.
"""

from __future__ import annotations

import re

from docx.text.paragraph import Paragraph

# XML 1.0 section 2.2 "Char" production (allowed): #x9 | #xA | #xD | [#x20-#xD7FF] | ...
# we remove anything in the disallowed ranges: 0x00-0x08, 0x0B, 0x0C, 0x0E-0x1F.
_INVALID = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F]")


def sanitize(value: str | None) -> str:
    if not value:
        return value or ""
    return _INVALID.sub("", value)


_orig_add_run = Paragraph.add_run


def _patched_add_run(self, text=None, style=None):  # type: ignore[no-untyped-def]
    return _orig_add_run(self, sanitize(text), style)


Paragraph.add_run = _patched_add_run  # type: ignore[method-assign]
