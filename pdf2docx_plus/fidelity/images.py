"""Restore inline (``wp:inline``) image emission and preserve alt-text.

Upstream `pdf2docx_plus._vendored.pdf2docx.common.docx.add_image`
unconditionally calls `add_floating_picture_pt`, which rewrites the
DOCX ``<w:drawing>`` tree into a ``<wp:anchor>`` with absolute page
coordinates.  That is correct for **floating** image blocks but wrong
for the much more common **inline** image span emitted from
``Image.make_docx``: a paragraph that should embed a logo or icon
flowing with text instead produces a floating overlay positioned at a
fixed (x, y) on the page.  Downstream consumers cannot reflow,
``align`` the image with its caption, or detect it as part of the
paragraph's text run.

This module:

1.  Replaces ``add_image`` with an inline-emission implementation that
    uses ``python-docx``'s native ``run.add_picture`` (yielding
    ``<wp:inline>``).  The bbox position arguments are ignored - the
    paragraph already encodes positioning via indent / alignment, so
    placing the image inline matches the line where upstream chose to
    emit it.

2.  Captures the PDF ``/Alt`` text for each image at extract time and
    stores it on the resulting raw dict.  ``Image.make_docx`` then
    propagates it onto the ``<wp:docPr>`` element so the DOCX picture
    carries a meaningful description (accessibility + search +
    diff-friendly).

3.  Leaves ``add_float_image`` (and the ``is_float_image_block`` branch
    of ``ImageBlock.make_docx``) untouched so genuinely-floating
    decorations - backdrops, watermark logos behind text - still emit
    as anchors.
"""

from __future__ import annotations

from io import BytesIO
from typing import Any

from docx.image.exceptions import UnrecognizedImageError  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from docx.shared import Pt  # type: ignore

import pdf2docx_plus._vendored.pdf2docx.common.docx as _docx
import pdf2docx_plus._vendored.pdf2docx.image.Image as _image_mod
import pdf2docx_plus._vendored.pdf2docx.image.ImagesExtractor as _imgext

from ..logging import get_logger

_log = get_logger("fidelity.images")


# ---------------------------------------------------------------------------
# 1. inline image emission
# ---------------------------------------------------------------------------


def _add_inline_image(paragraph: Any, image_stream: Any, width: float, height: float,
                      *, alt_text: str | None = None, alt_title: str | None = None) -> Any:
    """Insert an image into ``paragraph`` as an inline drawing.

    Returns the run that owns the picture, or ``None`` if the image
    bytes were unrecognised.  Caller is responsible for handling that
    case; ``Image.make_docx`` does so silently in upstream.
    """
    run = paragraph.add_run()
    try:
        run.add_picture(image_stream, width=Pt(width), height=Pt(height))
    except UnrecognizedImageError:
        _log.warning("unrecognised image; skipping")
        # remove the empty run we just added
        try:
            paragraph._p.remove(run._r)
        except Exception:
            pass
        return None

    if alt_text or alt_title:
        _set_alt_text(run, alt_title=alt_title, alt_text=alt_text)

    # restore single line spacing so the picture renders cleanly
    paragraph.paragraph_format.line_spacing = 1.0
    return run


def _set_alt_text(run: Any, *, alt_title: str | None, alt_text: str | None) -> None:
    """Set ``wp:docPr/@descr`` and ``wp:docPr/@title`` on the picture's drawing."""
    drawing = run._r.find(qn("w:drawing"))
    if drawing is None:
        return
    # walk to wp:docPr inside any wp:inline / wp:anchor
    for tag in ("wp:inline", "wp:anchor"):
        container = drawing.find(qn(tag))
        if container is not None:
            doc_pr = container.find(qn("wp:docPr"))
            if doc_pr is not None:
                if alt_text:
                    # OOXML uses 'descr'; many renderers accept 'title' for the short name.
                    doc_pr.set("descr", alt_text)
                if alt_title:
                    doc_pr.set("title", alt_title)
                return


def _patched_add_image(p: Any, image_path_or_stream: Any, x_pos: float, y_pos: float,
                       width: float, height: float) -> Any:
    """Drop-in replacement for upstream ``add_image``.

    Signature preserved (``x_pos`` / ``y_pos`` kept for compatibility with
    any downstream patch) but the implementation now emits ``<wp:inline>``
    via ``python-docx``'s built-in picture API instead of forcibly
    converting to ``<wp:anchor>``.

    Width/height passed by ``Image.make_docx`` are derived from the source
    bbox and therefore preserve the visual size of the image; the absolute
    (x_pos, y_pos) is intentionally discarded - in a real Word document
    images flow with text, not at absolute page coordinates.
    """
    # Try to recover an alt-text annotation that ``ImagesExtractor`` may
    # have stashed on the BytesIO wrapper or on a sidecar attribute of
    # the caller. ``Image.make_docx`` wraps ``self.image`` (raw bytes) in
    # a ``BytesIO``; ``_extract_alt_text`` looks for an ``_alt_text``
    # attribute placed there by the patched extractor.
    alt_text = getattr(image_path_or_stream, "_alt_text", None)
    alt_title = getattr(image_path_or_stream, "_alt_title", None)
    return _add_inline_image(p, image_path_or_stream, width, height,
                             alt_text=alt_text, alt_title=alt_title)


_docx.add_image = _patched_add_image


# ---------------------------------------------------------------------------
# 2. propagate alt-text from Image -> docx
# ---------------------------------------------------------------------------

# wrap Image.make_docx so we can pipe the alt-text stored on the Image
# instance through to the inline emit call. We only need to do this on
# the ``Image`` base class because both ImageBlock (inline branch) and
# ImageSpan go through the same code path.
_orig_image_make_docx = _image_mod.Image.make_docx


def _patched_image_make_docx(self, paragraph):  # type: ignore[no-untyped-def]
    stream = BytesIO(self.image)
    # smuggle alt-text onto the stream object so the patched add_image
    # can pull it without changing the upstream signature.
    alt_text = getattr(self, "_alt_text", None)
    alt_title = getattr(self, "_alt_title", None)
    if alt_text:
        stream._alt_text = alt_text  # type: ignore[attr-defined]
    if alt_title:
        stream._alt_title = alt_title  # type: ignore[attr-defined]
    _docx.add_image(
        paragraph,
        stream,
        self.bbox.x0,
        self.bbox.y0,
        self.bbox.x1 - self.bbox.x0,
        self.bbox.y1 - self.bbox.y0,
    )


_image_mod.Image.make_docx = _patched_image_make_docx  # type: ignore[method-assign]


# ---------------------------------------------------------------------------
# 3. read /Alt text out of the PDF and propagate into the raw image dict
# ---------------------------------------------------------------------------

def _extract_image_alt_text(page: Any, item: list) -> tuple[str, str] | tuple[None, None]:
    """Best-effort lookup of the ``/Alt`` field for an image on ``page``.

    PyMuPDF doesn't surface ``/Alt`` directly; we read the page's structure
    tree (``StructTreeRoot``) when present and match by the image's name
    (``item[7]``). On failure we silently return ``(None, None)`` so the
    image still emits without alt-text.
    """
    try:
        doc = page.parent
        img_name = item[7] if len(item) > 7 else None
        if not img_name:
            return (None, None)
        # Many PDFs store alt-text on a marked-content sequence; we scan
        # the page's content stream for a /Alt key in any ``/MCID``
        # property dict. This is a coarse but reliable heuristic.
        for xref in page.get_contents() or []:
            stream = doc.xref_stream(xref) or b""
            # Look for "/Alt (...)" pattern near "/Name /<img_name>"
            tag = f"/{img_name}".encode()
            idx = stream.find(tag)
            if idx < 0:
                continue
            window = stream[max(0, idx - 1024) : idx + 256]
            alt_start = window.find(b"/Alt(")
            if alt_start < 0:
                alt_start = window.find(b"/Alt (")
            if alt_start < 0:
                continue
            # take everything between the next '(' and the matching ')'
            paren = window.find(b"(", alt_start)
            if paren < 0:
                continue
            depth = 1
            i = paren + 1
            while i < len(window) and depth:
                c = window[i : i + 1]
                if c == b"\\":
                    i += 2
                    continue
                if c == b"(":
                    depth += 1
                elif c == b")":
                    depth -= 1
                    if depth == 0:
                        break
                i += 1
            if depth == 0:
                raw = window[paren + 1 : i]
                try:
                    alt = raw.decode("utf-8", errors="replace").strip()
                except Exception:
                    alt = ""
                if alt:
                    return (alt, img_name)
        return (None, None)
    except Exception as e:  # pragma: no cover - defensive
        _log.debug("alt-text extraction failed: %s", e)
        return (None, None)


# Wrap extract_images so the alt-text lands on the raw dict that
# eventually becomes an ``Image`` (or ImageBlock) instance.
_orig_extract_images = _imgext.ImagesExtractor.extract_images


def _patched_extract_images(self, *args, **kwargs):  # type: ignore[no-untyped-def]
    images = _orig_extract_images(self, *args, **kwargs)
    if not images:
        return images
    try:
        items = list(self._page.get_images(full=True))
        for raw_dict, item in zip(images, items):
            alt_text, alt_title = _extract_image_alt_text(self._page, item)
            if alt_text:
                raw_dict["alt_text"] = alt_text
            if alt_title:
                raw_dict["alt_title"] = alt_title
    except Exception as e:  # pragma: no cover - defensive
        _log.debug("alt-text propagation failed: %s", e)
    return images


_imgext.ImagesExtractor.extract_images = _patched_extract_images


# Also make sure the Image initializer keeps the alt-text fields when present.
_orig_image_init = _image_mod.Image.__init__


def _patched_image_init(self, raw=None):  # type: ignore[no-untyped-def]
    _orig_image_init(self, raw)
    if raw:
        alt = raw.get("alt_text")
        if alt:
            self._alt_text = alt
        title = raw.get("alt_title")
        if title:
            self._alt_title = title


_image_mod.Image.__init__ = _patched_image_init  # type: ignore[method-assign]
